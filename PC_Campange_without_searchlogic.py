import datetime as datetime
import ExcelInterface as EI
import InputConfigParser as ICF
import pygetwindow as pgw
import time
import KeyboardMouseSimulator as KMS
import logging
import threading
import sys
import re
import os
from docx import Document
import WordDocInterface as WDI
import Backlog_Handler as BH
import DCI_download_webinterface as DCIWB
import DCI_PC_Validation as DCIPC
import xlwings as xw
import PT_RR_Download_Fun_Reqs as PT
import shutil
date_time = datetime.datetime.now()

global fun


def progUpdate(func):
    global fun
    fun = func


def excel_popup(windowName):
    while (True):
        window_title = pgw.getActiveWindowTitle()
        global stop_threads
        excel_windows = pgw.getWindowsWithTitle("Excel")
        try:
            for each_excel_window in excel_windows:
                if windowName.split('.')[0] in each_excel_window.title:
                    each_excel_window.minimize()
                    each_excel_window.maximize()
                    if each_excel_window.isActive == False:
                        each_excel_window.activate()
                        break
                else:
                    each_excel_window.minimize()
        except:
            print("Exception in excel popup")
            break
        if pgw.getActiveWindowTitle() == "Microsoft Excel":
            time.sleep(1)
            KMS.rightArrow()
            time.sleep(1)
            KMS.pressEnter()
        active_window = pgw.getActiveWindow()
        if active_window is not None:
            active_window.minimize()
        if stop_threads:
            break


def get_values_from_campagne(Name, Start_date):
    projects = []
    Thématiques_inconnues, Thématiques_non_applicables = [], []
    NT_values, NA_values = [], []
    Type_of_Validation, Priority, Name_of_Project = '', '', ''
    logging.info("hi")
    try:
        path = ICF.getInputFolder() + "\\" + EI.findInputFiles()[19]
        print("path ---->", path)
        Campagnec_Book = EI.openExcel(path)
        Campagnec_Book.activate()
        Check_list_sTr_sheet = Campagnec_Book.sheets['Check-list Start']
        sheet_value = Check_list_sTr_sheet.used_range.value

        config_name = EI.searchDataInColCache(sheet_value, 1, 'Config')
        logging.info("config_name_1--------->", config_name)

        if not config_name.get('cellValue', []):
            # Increment the column index by 1
            new_column_index = 1 + 1
            config_name = EI.searchDataInColCache(sheet_value, new_column_index, 'Config')
            logging.info("config_name_2--------->", config_name)

        logging.info("config_name------------------>", config_name)
        Fun_name14 = EI.searchDataInColCache(sheet_value, 1, 'Writter of the campaign')
        logging.info("Fun_name14--------->", Fun_name14)
        if not Fun_name14.get('cellValue', []):
            new_column_index = 1 + 1
            Fun_name14 = EI.searchDataInColCache(sheet_value, new_column_index, 'Writter of the campaign')
            logging.info("Fun_name14--------->", Fun_name14)
        row, col = Fun_name14['cellPositions'][0]
        logging.info(row, col)
        EI.setDataInCell(Check_list_sTr_sheet, (row, col + 4), Name)
        EI.setDataInCell(Check_list_sTr_sheet, (row + 1, col + 4), Start_date)

        data1 = EI.searchDataInExcelCache(sheet_value, 'Result')
        roww, coll = data1['cellPositions'][0]
        logging.info("data1--------->", roww, coll)

        data2 = EI.searchDataInColCache(sheet_value, 1, 'Generated Campaign check')
        roww1, coll1 = data2['cellPositions'][0]
        Gen_camp_result = EI.getDataFromCell(Check_list_sTr_sheet, (roww1 + 1, coll))
        logging.info("Gen_camp_result------->", Gen_camp_result, roww1 + 1, coll)

        if Gen_camp_result == 'Previous campaign is published in docinfo':
            for i in range(1, 4):
                EI.setDataInCell(Check_list_sTr_sheet, (roww1 + i, coll + 1), 'Yes')
                logging.info("roww1+1+i, coll+1-->", roww1 + i, coll + 1)

        data3 = EI.searchDataInColCache(sheet_value, 1, 'Check the thematics of the campaign')
        roww2, coll2 = data3['cellPositions'][0]
        check_camp_them_result = EI.getDataFromCell(Check_list_sTr_sheet, (roww2 + 1, coll))
        logging.info("check_camp_them_result------->", check_camp_them_result)

        if check_camp_them_result == 'Check that all thematics are known':
            for i in range(1, 3):
                EI.setDataInCell(Check_list_sTr_sheet, (roww2 + i, coll + 1), 'Yes')
                logging.info("roww2 + i, coll+1------->", roww2 + i, coll + 1)

        for n, tuple_data in enumerate(config_name['cellPositions'][1:]):
            logging.info(tuple_data)
            row, col = tuple_data
            # row, col = config_name['cellPositions'][1]
            logging.info(row, col)
            col = col - 1
            Configurations = EI.getDataFromCell(Check_list_sTr_sheet, (row, col + 1))
            if Configurations:
                Type_of_Validation = EI.getDataFromCell(Check_list_sTr_sheet, (row, col + 4))
                Priority = EI.getDataFromCell(Check_list_sTr_sheet, (row, col + 5))
                Name_of_Project = EI.getDataFromCell(Check_list_sTr_sheet, (row, col + 6))
                logging.info("Configurations->, Type_of_Validation-------->", Configurations, Type_of_Validation)
            if Configurations is None:
                logging.info("ooooopppppppp")
                logging.info("row, colrow, col----->", row, col)
                try:
                    Configurations = EI.getDataFromCell(Check_list_sTr_sheet, (row, col))
                    Type_of_Validation = EI.getDataFromCell(Check_list_sTr_sheet, (row, col + 5))
                    Priority = EI.getDataFromCell(Check_list_sTr_sheet, (row, col + 6))
                    Name_of_Project = EI.getDataFromCell(Check_list_sTr_sheet, (row, col + 7))
                    logging.info("row, col--------->", row, col)
                    logging.info("Configurations ifConfigurations---------->", Configurations)
                    col = int(col) - 1
                    logging.info("col try---->", col)
                except:
                    Configurations = EI.getDataFromCell(Check_list_sTr_sheet, (row, col + 2))
                    logging.info("row, col+2)--------->", row, col + 2)
                    Type_of_Validation = EI.getDataFromCell(Check_list_sTr_sheet, (row, col + 5))
                    Priority = EI.getDataFromCell(Check_list_sTr_sheet, (row, col + 6))
                    Name_of_Project = EI.getDataFromCell(Check_list_sTr_sheet, (row, col + 7))
                    logging.info("Configurations except Configurations---------->", Configurations)
                    col = int(col + 2) - 1
                    logging.info("col except---->", col)

            print("DFGSFDAFGAERF------->", Configurations, Type_of_Validation, Priority, Name_of_Project)

            if str(n + 1) == str(1):
                # EI.setDataFromCell(Check_list_sTr_sheet, (row, col + 4), 'PC1')
                Config_sheet = 'Campagne Config 1'
                logging.info("Config_sheet1-------->", Config_sheet)
                logging.info("row, colrow, colrow, col------>", row, col)
                project1 = EI.getDataFromCell(Check_list_sTr_sheet, (row, col + 6))
                logging.info("project-------->", project1)
                projects.append(project1.split('(')[0].strip())
                Thématiques_inconnue, Thématiques_non_applicable, NT_value, NA_value = DCIPC.getthematics(Config_sheet)
                NT_values.append(NT_value)
                NA_values.append(NA_value)
                Thématiques_inconnues.append(Thématiques_inconnue)
                Thématiques_non_applicables.append(Thématiques_non_applicable)

            else:
                n = n - 1
                # EI.setDataFromCell(Check_list_sTr_sheet, (row, col + 4), 'Compl. PC'+str(n + 1))
                Config_sheet = 'Campagne Config ' + str((n + 1) + 1)
                logging.info("Config_sheet2-------->", Config_sheet)
                project2 = EI.getDataFromCell(Check_list_sTr_sheet, (row, col + 6))
                Thématiques_inconnue, Thématiques_non_applicable, NT_value, NA_value = DCIPC.getthematics(Config_sheet)
                NT_values.append(NT_value)
                NA_values.append(NA_value)
                Thématiques_inconnues.append(Thématiques_inconnue)
                Thématiques_non_applicables.append(Thématiques_non_applicable)
                projects.append(project2.split('(')[0].strip())

        logging.info("projects.append(project1)---------->", projects)
        logging.info("1stsheet--------->", Thématiques_inconnues, Thématiques_non_applicables)
        logging.info("NT_values and NA_values--------->", NT_values, NA_values)

        # List of sheets to hide
        sheets_to_hide = ["Sommaire", "Read_IT"]

        # Iterate through sheets and hide the specified ones
        for sheet_name in sheets_to_hide:
            try:
                sheet = Campagnec_Book.sheets[sheet_name]
                sheet.api.Visible = False  # Set Visible property to hide the sheet
            except KeyError:
                print(f"Sheet '{sheet_name}' not found.")

        Campagnec_Book.save()
        fun(10)
    except Exception as ex:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        exp_fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
        print(f"{ex} line no. {exc_tb.tb_lineno} file name: {exp_fname}")
    return projects, Thématiques_inconnues, Thématiques_non_applicables, NT_values, NA_values


# Function to check if a sheet is present in an Excel file
def is_sheet_present(file_path, sheet_name, filename):
    try:
        global stop_threads
        stop_threads = False
        t1 = threading.Thread(target=excel_popup, args=(filename,))
        t1.start()
        # Open the Excel file
        workbook = xw.Book(file_path)
        stop_threads = True

        # Check if the sheet is present
        if any(sheet_name.strip().lower() == sheet.name.strip().lower() for sheet in workbook.sheets):
            return True
        else:
            workbook.close()
            return False
    except Exception as e:
        stop_threads = True
        print(f"Error: {e}")
        return False


def to_find_sheet_inPLM(desired_sheet_name):
    sheet_found_flag = 0
    file_name = ''
    input_folder = ICF.getInputFolder()
    # List of file extensions to filter Excel files
    file_extensions = [".xlsx", ".xlsm"]
    # Sheet name to check for
    # Flag to track whether the sheet is found in any file
    sheet_found = False

    # Loop through each file in the folder
    for file_name in os.listdir(input_folder):
        # if any(file_name.endswith(ext) and "PLM" in file_name for ext in file_extensions):
        if any(file_name.endswith(ext) for ext in file_extensions):
            # if file_name.endswith(file_extension) and "PLM" in file_name:
            file_path = os.path.join(input_folder, file_name)

            # Check if the desired sheet is present in the current Excel file
            if is_sheet_present(file_path, desired_sheet_name, file_name):
                sheet_found_flag = 1
                print(f"The sheet '{desired_sheet_name}' is present in the file: {file_name}")
                # Set the flag to True and break out of the loop
                sheet_found = True
                break
    # Check if the sheet was not found in any file
    if not sheet_found:
        print(f"The sheet '{desired_sheet_name}' is not present in any PLM files.")
    return sheet_found_flag, file_name


def checkinplm(result_tuples, file_name, NT_value, NA_value, desired_sheet_name, sheet_found_flag, Thématiques_inconnues, desired_architecture):
    logging.info("desired_sheet_name--------------->", desired_sheet_name)
    NA_content = []
    path = ICF.getInputFolder() + "\\" + file_name
    logging.info("path ---->", path)
    global stop_threads
    stop_threads = False
    t1 = threading.Thread(target=excel_popup, args=(file_name,))
    t1.start()
    try:
      PLM_EE_Book = EI.openExcel(path)
    except:
        stop_threads = True
    stop_threads = True
    PLM_EE_Book.activate()
    try:
        if sheet_found_flag:
            func_sheet = PLM_EE_Book.sheets[desired_sheet_name]
            # Convert the list to an Excel range object
            sheet_range = func_sheet.range('A1').expand()
            # Find the last non-empty row
            maxrow = sheet_range.end('up').row
            logging.info("maxrow--------->", maxrow)
            sheet_value = func_sheet.used_range.value
            Fun_name9 = EI.searchDataInColCache(sheet_value, 7, 'Name')
            logging.info("Fun_name9----->", Fun_name9)
            # Check if 'cellPositions' is not empty before accessing its elements
            if Fun_name9['cellPositions']:
                row, col = Fun_name9['cellPositions'][0]
                time.sleep(2)
                logging.info(row, col)
                for i in result_tuples:
                    if i[1] == '--' or i[1] == 'X':
                        logging.info("i--------->", i[0])
                        result_tuple = i[0]
                        logging.info("result_tuple.split('_')[0]---->", result_tuple.split('_')[0])
                        try:
                            PLM_sheet_Valeur = EI.searchDataInColCache(sheet_value, col, result_tuple.split('_')[0])
                            logging.info("PLM_sheet_Valeur--------->", PLM_sheet_Valeur)
                            # Condition 0
                            if PLM_sheet_Valeur['count'] == 0:
                                logging.info("PLM_sheet_Valeur['count'] == 0")
                                b = result_tuple + ':  According to the DecliEE and Silhouette files, this thematic is NA'
                                NA_content.append(b)

                            elif PLM_sheet_Valeur['count'] == 1:
                                logging.info("PLM_sheet_Valeur['count'] == 1")
                                row, col = PLM_sheet_Valeur['cellPositions'][0]
                                thematic_value = EI.getDataFromCell(func_sheet, (row, col + 2))
                                # Condition 1
                                if i[1] == '--':
                                    #  -- !=
                                    if result_tuple != thematic_value:
                                        b = result_tuple + ':  According to the DecliEE and Silhouette files, this thematic is NA, but is fixed on the value ' + thematic_value
                                        NA_content.append(b)
                                    else:
                                        b = result_tuple + ':  Raise QIA of Silhouette.'
                                        NA_content.append(b)
                                # Condition 2
                                if i[1] == 'X':
                                    if result_tuple != thematic_value or result_tuple == thematic_value:
                                        b = result_tuple + ':  Raise QIA of Silhouette.'
                                        NA_content.append(b)

                            elif PLM_sheet_Valeur['count'] >= 2:
                                logging.info("PLM_sheet_Valeur['count'] > 2")
                                result_list = []
                                all_thematic_values = []
                                if i[1] == '--':
                                    for cellposition in PLM_sheet_Valeur['cellPositions']:
                                        row, col = cellposition
                                        logging.info("row,col", row, col)
                                        thematic_value = EI.getDataFromCell(func_sheet, (row, col + 2))
                                        all_thematic_values.append(thematic_value)
                                    print("all_thematic_values---->", all_thematic_values)

                                    if result_tuple in all_thematic_values:
                                        b = result_tuple + ':  Raise QIA of Silhouette.'
                                        NA_content.append(b)
                                    else:
                                        result = "|".join(all_thematic_values)
                                        b = result_tuple + ':  According to the DecliEE and Silhouette files, this thematic is NA, but is optional on the value ' + result
                                        NA_content.append(b)
                                if i[1] == 'X':
                                    for cellposition in PLM_sheet_Valeur['cellPositions']:
                                        row, col = cellposition
                                        logging.info("row,col", row, col)
                                        thematic_value = EI.getDataFromCell(func_sheet, (row, col + 2))
                                        if result_tuple != thematic_value or result_tuple == thematic_value:
                                            b = result_tuple + ':  Raise QIA of Silhouette.'
                                            NA_content.append(b)
                                            break
                        except Exception as ex:
                            exc_type, exc_obj, exc_tb = sys.exc_info()
                            exp_fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
                            print(f"{ex} line no. {exc_tb.tb_lineno} file name: {exp_fname}")
                            print("Thematics are not present in the PLM_Sheet.")
                    if i[1] == 'opt':
                        logging.info("i--------->", i[0])
                        result_tuple = i[0]
                        logging.info("result_tuple.split('_')[0]---->", result_tuple.split('_')[0])
                        try:
                            PLM_sheet_Valeur = EI.searchDataInColCache(sheet_value, col, result_tuple.split('_')[0])
                            logging.info("PLM_sheet_Valeur--------->", PLM_sheet_Valeur)
                            if PLM_sheet_Valeur['count'] > 0:
                                logging.info("PLM_sheet_Valeur['count'] > 0")
                                result_list = []
                                for cellposition in PLM_sheet_Valeur['cellPositions']:
                                    row, col = cellposition
                                    logging.info("row,col", row, col)
                                    thematic_value = EI.getDataFromCell(func_sheet, (row, col + 2))
                                    if result_tuple != thematic_value:
                                        b = result_tuple + ':  Raise QIA of Silhouette.'
                                        NA_content.append(b)
                                        break
                        except Exception as ex:
                            exc_type, exc_obj, exc_tb = sys.exc_info()
                            exp_fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
                            print(f"{ex} line no. {exc_tb.tb_lineno} file name: {exp_fname}")
                            print("Thematics are not present in the PLM_Sheet.")
    except:
        print("sheet are not present in the PLM_Sheet.")

    doc_content = ['Total number of NT files: ' + str(NT_value), 'Total number of NA files: ' + str(NA_value),
                   'NA because of following thematiques: ']
    print("desired_architecture--desired_architecture----desired_architecture----------->", desired_architecture)
    # try:
    # Thématiques_inconnues_COMMENT = ''
    Thématiques_inconnues_COMMENT = []
    if Thématiques_inconnues:
        if desired_architecture == 'NEA R1.1' or desired_architecture == 'NEA R1.2':
            version = 'NEA R1'
            Thématiques_inconnues_COMMENT = [
                'The following thematics:  ' + Thématiques_inconnues + ' are only applicable for ' + version + ' according to the CONFIG_THEMATIQUES file.']
        if desired_architecture == 'NEA R1' or desired_architecture == 'NEA R1.0':
            version = 'NEA R1.x'
            Thématiques_inconnues_COMMENT = [
                'The following thematics:  ' + Thématiques_inconnues + ' are only applicable for ' + version + ' according to the CONFIG_THEMATIQUES file.']
    NA_NT_content = doc_content + NA_content + Thématiques_inconnues_COMMENT
    # NA_NT_content = NA_content
    logging.info("NA_NT_content-------->", NA_NT_content)
    PLM_EE_Book.close()
    # except Exception as ex:
    #     exc_type, exc_obj, exc_tb = sys.exc_info()
    #     exp_fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
    #     print(f"{ex} line no. {exc_tb.tb_lineno} file name: {exp_fname}")

    return NA_NT_content


def append_to_word_document(contents, thematics_not_applicable_for_Feps, output_path, docm):
    # Iterate through the content
    # for content in contents:
    # Assume the first element is the heading
    heading = contents[0]

    # Check if the heading is a list of strings
    if isinstance(heading, list):
        # Join the list of strings into a single string
        heading = ' '.join(heading)

    # Add the heading to the document with bold formatting
    docm.add_paragraph(heading.strip(), style='Heading1')

    # Iterate through the rest of the content
    for line in contents[1:]:
        if isinstance(line, list):
            # Treat '\n' as a new line in the Word document
            if '\n' in line:
                docm.add_paragraph()
        else:
            sanitized_line = ''.join(c for c in line if c.isprintable())

            # Check if the line is not empty after sanitization
            if sanitized_line.strip():
                docm.add_paragraph(sanitized_line.strip())

    if thematics_not_applicable_for_Feps:
        thematics_not_applicable_for_Feps = [thematics_not_applicable_for_Feps[0]] + [item for sublist in thematics_not_applicable_for_Feps[1:] for item in sublist]
        print("append_to_word_document_thematics_not_applicable_for_Feps----------->", thematics_not_applicable_for_Feps)
        print("table funcion inside-------->", thematics_not_applicable_for_Feps)
        # for table add in the word document
        table_doc = docm.add_table(rows=len(thematics_not_applicable_for_Feps),
                                   cols=len(thematics_not_applicable_for_Feps[0]))
        for i, row in enumerate(thematics_not_applicable_for_Feps):
            for j, cell in enumerate(row):
                table_doc.cell(i, j).text = cell
        table_doc.style = 'Table Grid'
    docm.save(output_path)


def findlines(them_line, Thématiques_non_applicable_codes, Thématiques_inconnues_codes, Feps):
    # print("them_line, Thématiques_non_applicable_codes_result---------->",them_line, Thématiques_non_applicable_codes)
    combined_list = Thématiques_non_applicable_codes + Thématiques_inconnues_codes

    logging.info("combined_list------>", combined_list)
    # Split the multiline string into a list of lines
    lines = them_line.split('\n')

    # Remove lines containing codes from Thématiques_non_applicable_codes_result
    filtered_lines = [line for line in lines if not any(code in line for code in combined_list)]
    logging.info("filtered_lines---------------filtered_lines, len(filtered_lines), len(lines)->", filtered_lines, len(filtered_lines), len(lines))
    # Inform the user if lines were removed
    if len(filtered_lines) != 0:
        b = 'Please find the supporting requirement for ' + Feps
        # print(f"Please find the supporting requirement for {Feps}.")
        print(f"Please find the supporting requirement for {Feps} or req is not present in the Searchlogic doc.")
    else:
        b = Feps + ' is not applicable to ARCH or Project.'
        logging.info("Feps is not applicable to ARCH or Project.")
    # Join the filtered lines back into a multiline string
    result = '\n'.join(filtered_lines)
    # print(result)
    content = b + result
    return content, len(filtered_lines)


def interface_findlines(them_line, Thématiques_non_applicable_codes, Thématiques_inconnues_codes, Feps):
    # print("them_line, Thématiques_non_applicable_codes_result---------->",them_line, Thématiques_non_applicable_codes)
    combined_list = Thématiques_non_applicable_codes + Thématiques_inconnues_codes

    print("combined_list------>", combined_list)
    # Split the multiline string into a list of lines
    lines = them_line.split('\n')

    # Remove lines containing codes from Thématiques_non_applicable_codes_result
    filtered_lines = [line for line in lines if not any(code in line for code in combined_list)]
    print("filtered_lines---------------filtered_lines, len(filtered_lines), len(lines)->", filtered_lines, len(filtered_lines), len(lines))
    # Inform the user if lines were removed
    if len(filtered_lines) != 0:
        b = 'Please find the requirement signal for ' + Feps
        print(f"Please find the requirement signal for {Feps}.")
    else:
        b = Feps + ' is not applicable to ARCH or Project.'
        print("Feps is not applicable to ARCH or Project.")

    # Join the filtered lines back into a multiline string
    result = '\n'.join(filtered_lines)
    # print(result)
    content = b + result
    return content, len(filtered_lines)


def projectcode(desired_architecture, project_name):
    logging.info("desired_architecture, project_name---->", desired_architecture, project_name)
    project_code = ''
    # Split the strings by underscores and spaces
    architecture_parts = desired_architecture.split('_')
    project_parts = project_name.split('_')
    logging.info("architecture_parts----project_parts------>", architecture_parts, project_parts)

    # Find the common part (project code) by checking intersections
    common_parts = set(architecture_parts) & set(project_parts)
    logging.info("common_parts---------->", common_parts)
    # If there is a common part, extract it
    if common_parts:
        project_code = list(common_parts)[0]
        logging.info("Project Code:", project_code)
    else:
        print("Project code No common part found.")
    return project_code


def extract_unique_elements(data, element_index):
    # Extract the specified element from each tuple
    result = [t[element_index] for t in data]

    # Remove duplicates by converting the list to a set and then back to a list
    unique_result = list(set(result))

    # Convert the list to a comma-separated string
    return ",".join(unique_result)


def get_fun_reqs(functional_requirements):
    table_lists, Doc_lists = [], []
    PT.getdocs(functional_requirements, reqFound=[])
    req_present_in_download_docs = []
    for filename in os.listdir(ICF.getInputFolder()):
        if filename.endswith(('.doc', '.docx')):
            # Process the file (replace this with your logic)
            print(f"Processing::: {filename}")
            try:
                WDI.save_as_docx(ICF.getInputFolder() + "\\" + filename)
            except:
                pass
            Doc = os.path.join(ICF.getInputFolder(), filename)
            print("Keyword_file_path----->", Doc)
            # for rv in all_filtered_req_list[:]:  # Iterate over a copy of the list
            for rv in functional_requirements:
                ReqName, ReqVer = PT.getReqVer(rv)
                print("reqName, reqVer---->", ReqName, ReqVer)
                try:
                    RqTable, Content, Doc = WDI.getContent(Doc, ReqName, ReqVer)
                    table_lists.append(RqTable)
                    Doc_lists.append(Doc)
                    if Content != None and Content != -1 and Content != '':
                        req_present_in_download_docs.append(rv)
                        functional_requirements.remove(rv)  # Remove the element from the list
                except Exception as ex:
                    exc_type, exc_obj, exc_tb = sys.exc_info()
                    exp_fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
                    print(
                        f"\nSomething went wrong getting table {ex} line no. {exc_tb.tb_lineno} file name: {exp_fname}")

            print("req_present_in_download_docs------------>", req_present_in_download_docs)
            print("rv_in_download_docs------------>", functional_requirements)
            # Remove elements present in req_present_in_download_docs from all_filtered_req_list
            functional_requirements = [item for item in functional_requirements if item not in req_present_in_download_docs]

    print("before_table_lists--------->", table_lists)
    print("before_Doc_lists--------->", Doc_lists)

    result_tuples = list(zip(table_lists, Doc_lists))
    print("before_result_tuples---->", result_tuples)
    table_lists = [(a, b) for (a, b) in result_tuples if a != '' and a != -1]
    print("before_result_tuples---->", table_lists)
    for data in table_lists:
        PT.addDataInDocument(data, "T")
    print("Processing complete.")


def functionalreq(Feps, functional_requirements, Thématiques_inconnue, Thématiques_non_applicable, desired_architecture, project_code, full_architecture, Req_impacted_sheet, cleaned_requirements, Impacted_sheets):
    impacted_sheets = Impacted_sheets.split('\n')
    Impacted_sheets = ','.join(impacted_sheets)

    Thématiques_inconnues_codes = []
    Thématiques_non_applicable_codes = []

    Thématiques_inconnues_flags = []
    Thématiques_non_applicables_flags = []
    Functionalreqscontent = []
    Thématiques_non_applicable_codes_result, Thématiques_inconnues_codes_result = '', ''
    req_bal_list = []
    all_rows_list = []
    global iteration_project
    global iteration_project_index
    for req in functional_requirements:
        # Split the Thématiques strings into lists
        them_code, them_line = findReqinSearchLogicDoc(req)
        if them_code == -1 and them_line == -1:
            them_code, them_line = findReqinC4Doc(req)
            if them_code == -1 and them_line == -1:
                get_fun_reqs(functional_requirements)
                # In the below line it will create the C4 doc in the output folder
                them_code, them_line = findReqinC4Doc(req)
            # if file_path:
            #     print("if part")
            #     them_code, them_line = findReqinC4Doc(req)
            # else:
            #     print("else part")
            #     get_fun_reqs(functional_requirements)
            #     # In the below line it will create the C4 doc in the output folder
            #     them_code, them_line = findReqinC4Doc(req)
                #req_bal_list.append(req)
        if them_code == -1 and them_line == -1:
            req_bal_list.append(req)
            continue

        # inconnues_list = Thématiques_inconnue.split('|')
        # non_applicables_list = Thématiques_non_applicable.split('|')
        print("Thématiques_inconnue-------------->", Thématiques_inconnue, Thématiques_non_applicable)
        if Thématiques_inconnue is not None:
            inconnues_list = Thématiques_inconnue.split('|')
        else:
            inconnues_list = []

        if Thématiques_non_applicable is not None:
            non_applicables_list = Thématiques_non_applicable.split('|')
        else:
            non_applicables_list = []
        them_line_list = them_code
        code = ''
        # Check if any code in them_line is present in Thématiques_inconnues or Thématiques_non_applicables
        for code in them_line_list:
            if code in inconnues_list:
                inconnues_list_flag = 1
                # print(f"Code {code} is present in Thématiques_inconnues.")
                Thématiques_inconnues_codes.append(code)
                Thématiques_inconnues_flags.append(inconnues_list_flag)

            else:
                inconnues_list_flag = 2
                # print(f"Code {code} is not present in Thématiques_inconnues.")
                Thématiques_inconnues_flags.append(inconnues_list_flag)

            if code in non_applicables_list:
                non_applicables_list_flag = 1
                # print(f"Code {code} is present in Thématiques_non_applicables.")
                Thématiques_non_applicable_codes.append(code)
                Thématiques_non_applicables_flags.append(non_applicables_list_flag)

            else:
                non_applicables_list_flag = 2
                # print(f"Code {code} is not present in Thématiques_non_applicables.")
                Thématiques_non_applicables_flags.append(non_applicables_list_flag)

        Thématiques_inconnue_code_result = ','.join(Thématiques_inconnues_codes)
        logging.info("Thématiques_inconnues_codes_result--->", Thématiques_inconnue_code_result)
        # Step 1: Convert the string to a list
        codes_list = Thématiques_inconnue_code_result.split(',')

        # Step 2: Convert the list to a set to remove duplicates
        unique_codes_set = set(codes_list)

        # Step 3: Convert the set back to a string
        Thématiques_inconnues_codes_result = ','.join(unique_codes_set)
        logging.info("After removing duplicates:", Thématiques_inconnues_codes_result)

        Thématiques_non_applicable_code_result = ','.join(Thématiques_non_applicable_codes)
        logging.info("Thématiques_non_applicable_codes_result--->", Thématiques_non_applicable_code_result)
        # Step 1: Convert the string to a list
        codes_list = Thématiques_non_applicable_code_result.split(',')

        # Step 2: Convert the list to a set to remove duplicates
        unique_codes_set = set(codes_list)

        # Step 3: Convert the set back to a string
        Thématiques_non_applicable_codes_result = ','.join(unique_codes_set)
        logging.info("After removing duplicates:", Thématiques_non_applicable_codes_result)

        content, filtered_lines = findlines(them_line, Thématiques_non_applicable_codes_result,
                                            Thématiques_inconnues_codes_result, Feps)
        logging.info("Thematics lines applicable to project or Arch from the functional requirement------>", content,
                     filtered_lines)
        Functionalreqs = content + ' length of the applicable thematcis lines ' + str(filtered_lines) + ' ' + req
        Functionalreqscontent.append(Functionalreqs)

    logging.info("Functionalreqscontent----->", Functionalreqscontent)
    master_list = list(set(Functionalreqscontent))
    print("master_list--------->", master_list)
    reqlist = []
    notreqlist = []
    suplist = []
    for element in master_list:
        if 'length of the applicable thematcis lines 0' in element:
            parts = element.split("length of the applicable thematcis lines 0")
            if len(parts) > 1:
                reqlist.append("'" + parts[1].strip() + "'")
        if 'the supporting requirement' in element:
            parts = element.split("length of the applicable thematcis lines ")
            if len(parts) > 1:
                getreq = "'" + parts[1].strip() + "'"
                req = getreq.split(" ")[1]
                suplist.append(req)
        else:
            parts = element.split("length of the applicable thematcis lines ")
            if len(parts) > 1:
                getreq = "'" + parts[1].strip() + "'"
                req = getreq.split(" ")[1]
                notreqlist.append(req)

    not_applicable_fun_req = ','.join(reqlist)
    len_not_applicable_fun_req = len(reqlist)
    logging.info("not_applicable_fun_req- len_not_applicable_fun_req------>", not_applicable_fun_req, len_not_applicable_fun_req)

    applicable_fun_req = ','.join(notreqlist)
    len_applicable_fun_req = len(notreqlist)
    logging.info("applicable_fun_req---len_applicable_fun_req---->", applicable_fun_req, len_applicable_fun_req)

    applicable_sup_req = ','.join(suplist)
    len_applicable_sup_req = len(suplist)
    logging.info("applicable_sup_req---len_applicable_sup_req---->", applicable_sup_req, len_applicable_sup_req)

    if len_not_applicable_fun_req and len_applicable_fun_req == 0:
        third_row_in_table = 'For the ' + Feps + ' have the requirement ' + not_applicable_fun_req + ' which are NA for our project ' + project_code + '. Because of thematics.'
        logging.info("third_row_in_table for functional reqs_not_applicable---------->", third_row_in_table)
        fourth_row_in_table = 'When thematic are applicable for our project this  ' + Feps + '  can be treated.'
        logging.info("modified_arch_for_table----------->", full_architecture)
        modified_arch_for_table = full_architecture.replace(" ", "_")
        fifth_row_in_table = "Cannot be tested because the thematics are not applicable for " + modified_arch_for_table
        second_row_in_table = "No"
        first_row_in_table = Feps
        # Create a list and append the strings
        rows_list = [first_row_in_table, second_row_in_table, third_row_in_table, fourth_row_in_table,
                     fifth_row_in_table]
        # Print the resulting list
        logging.info(rows_list)
        all_rows_list.append(rows_list)

    if len_applicable_sup_req:
        print("functional req Impacted_sheets------------>", Impacted_sheets)
        if iteration_project[iteration_project_index]==True:
            # PT.main(Impacted_sheets)
            iteration_project[iteration_project_index] = False
            print("first_iteration----------->")
        else:
            print("RUNING AFTER JSKDFJHK")
            pass
        third_row_in_table = 'For the ' + Feps + ' have the requirement ' + applicable_sup_req + ' Need supporting Requirements or it may not present in the Search logic folder.'
        logging.info("third_row_in_table for functional Sup_reqs---------->", third_row_in_table)
        fourth_row_in_table = '--'
        fifth_row_in_table = "--"
        second_row_in_table = "No"
        first_row_in_table = Feps
        # Create a list and append the strings
        rows_list = [first_row_in_table, second_row_in_table, third_row_in_table, fourth_row_in_table,
                     fifth_row_in_table]
        # Print the resulting list
        print(rows_list)
        all_rows_list.append(rows_list)

    if not_applicable_fun_req:
        third_row_in_table = 'For the ' + Feps + ' have the requirement ' + not_applicable_fun_req + ' having the thematic ' + Thématiques_non_applicable_codes_result, Thématiques_inconnues_codes_result + ' which is NA for our project ' + project_code
        fourth_row_in_table = 'When thematic are applicable for our project this  ' + Feps + '  can be treated.'
        print("modified_arch_for_table----------->", full_architecture)
        modified_arch_for_table = full_architecture.replace(" ", "_")
        fifth_row_in_table = "Cannot be tested because the thematics are not applicable for " + modified_arch_for_table
        second_row_in_table = "No"
        first_row_in_table = Feps
        # Create a list and append the strings
        rows_list = [first_row_in_table, second_row_in_table, third_row_in_table, fourth_row_in_table,
                     fifth_row_in_table]
        # Print the resulting list
        print(rows_list)
        all_rows_list.append(rows_list)
    return all_rows_list


def impacted_sheet(dcireqs_applicable, Req_impacted_sheet):
    # Remove leading and trailing single quotes, split the string by comma to get individual items
    dcireqs_applicable_items = dcireqs_applicable.strip("'").split("', '")

    # Initialize the list to store the matching second elements
    matching_second_elements = []

    # Iterate through each item in dcireqs_applicable_items
    for dcireq in dcireqs_applicable_items:
        # Iterate through each tuple in Req_impacted_sheet
        for item in Req_impacted_sheet:
            # Check if the first element of the tuple matches dcireq
            if item[0] == dcireq:
                # If a match is found, retrieve the second element of the tuple
                matching_second_elements.append(item[1])

    # Convert the list of matching second elements to a comma-separated string
    matching_second_elements_str = ', '.join(matching_second_elements)

    # print("Matching elements:", matching_second_elements_str)
    return matching_second_elements_str


def interfacereq(Feps, interface_requirements, Thématiques_inconnue, Thématiques_non_applicable, excluded_sheets,
                 desired_architecture, project_code, full_architecture, Req_impacted_sheet, all_DCI_Reqs_sheets,
                 Impacted_sheets):
    impacted_sheets = Impacted_sheets.split('\n')
    Impacted_sheets = ','.join(impacted_sheets)

    Thématiques_inconnues_codes, Thématiques_non_applicable_codes = [], []
    Thématiques_inconnues_flags, Thématiques_non_applicables_flags = [], []
    master_list, all_rows_list = [], []
    dci_req_results, dci_req_not_found_in_file, dci_req_found_in_files, DCI_req_Arch_not_match = [], [], [], []
    Thématiques_non_applicable_codes_result, Thématiques_inconnues_codes_result, col_to_get_NT_NA_values = '', '', ''
    dcireq_not_applicable, signal_present_sheetss, signal_present_sheetsss = [], [], []
    NA_values_for_signal, NA_DCI_signal, NA_values_for_signal_sheet, NA_values_for_req = [], [], [], []
    NT_values_for_signal, NT_DCI_signal, NT_values_for_signal_sheet, NT_values_for_req = [], [], [], []

    # ssd_folder = ICF.getSsdFolder()
    ssd_folder = ICF.getDicFolder()

    if not os.path.exists(ssd_folder) or not os.listdir(ssd_folder):
        # Either the directory doesn't exist or it's empty
        # Handle the condition here
        logging.info("Directory is either not present or empty.")
        TestPlan = EI.findInputFiles()[1]
        print("PT---->", TestPlan)
        tpBook = EI.openExcel(ICF.getInputFolder() + "\\" + TestPlan)
        DCIWB.download_documents(tpBook, ('xls',), ('dci'), True)

    # checking the interface req is present in the functional DCI files
    for dci_Req in interface_requirements:
        logging.info("dci--->", dci_Req)
        # for i in os.listdir(ICF.getSsdFolder()):
        for i in os.listdir(ICF.getDicFolder()):
            if ('dci' in i.lower().strip()):
                try:
                    # dciBook = EI.openExcel(ICF.getSsdFolder() + "\\" + i)
                    dciBook = EI.openExcel(ICF.getDicFolder() + "\\" + i)
                    logging.info("dciBookdciBook------------------->", dciBook)
                    dciInfo = DCIPC.getDciInfo(dciBook, dci_Req)
                    logging.info("DCI req-->", dciInfo)

                    # Check if DCI_req1 has empty values for specified keys
                    if any(dciInfo[key] for key in ['dciSignal', 'arch', 'thm', 'dciReq']):
                        logging.info("DCI req1 dictionary is not empty")
                        req_architecture = DCIPC.dciArch(dciInfo)
                        modified_architecture = desired_architecture.replace(" ", "")
                        logging.info("modified_architecture-req_architecture---->", modified_architecture, req_architecture)
                        # Check if all conditions are true
                        if (req_architecture == modified_architecture or req_architecture == "NEAR1.0" and modified_architecture == "NEAR1"
                                or req_architecture == "NEAR1" and modified_architecture == "NEAR1.0"
                                or req_architecture == "NEAR1.1" and modified_architecture == "NEAR2"
                                or req_architecture == "NEAR2" and modified_architecture == "NEAR1.1"
                                or req_architecture == "NEAR1.2" and modified_architecture == "NEAR3"
                                or req_architecture == "NEAR3" and modified_architecture == "NEAR1.2"):
                            logging.info("All conditions are true")
                            print("Arch matched with user and Functional DCI file", req_architecture, modified_architecture)
                            dci_req_results.append(dciInfo)
                            dci_req_found_in_files.append(dci_Req)
                            break
                        else:
                            print(f"Arch is not match with the {dci_Req}.")
                            DCI_req_Arch_not_match.append(dci_Req)
                    dciBook.close()
                except Exception as ex:
                    exc_type, exc_obj, exc_tb = sys.exc_info()
                    exp_fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
                    logging.info(f"DCI Req is not present in the DCI files present in the input folder.")
                    dci_req_not_found_in_file.append(dci_Req)

    DCI_req_Arch_not_match = list(set(DCI_req_Arch_not_match))
    print("dci_req_results------->", dci_req_results)
    print("dci_req_found_in_files------->", dci_req_found_in_files)
    print("DCI_req_Arch_not_match------->", DCI_req_Arch_not_match)

    # checking the thematics of the req with campagne thematics
    for result in dci_req_results:
        dciSignal = result['dciSignal']
        arch = result['arch']
        thm = result['thm']
        dciReq = result['dciReq']
        logging.info("thm------------->", dciReq, thm)
        thematic_string = thm.decode('utf-8')
        logging.info("thematic_string-------->", thematic_string)
        them_code, them_line = DCIPC.getThematiccode_for_DCI(thematic_string)
        try:
            print("Thématiques_inconnue-------------->", Thématiques_inconnue, Thématiques_non_applicable)
            if Thématiques_inconnue is not None:
                inconnues_list = Thématiques_inconnue.split('|')
            else:
                inconnues_list = []

            if Thématiques_non_applicable is not None:
                non_applicables_list = Thématiques_non_applicable.split('|')
            else:
                non_applicables_list = []

            them_line_list = them_code
            # Check if any code in them_line is present in Thématiques_inconnues or Thématiques_non_applicables
            for code in them_line_list:
                if code in inconnues_list:
                    inconnues_list_flag = 1
                    # print(f"Code {code} is present in Thématiques_inconnues.")
                    Thématiques_inconnues_codes.append(code)
                    Thématiques_inconnues_flags.append(inconnues_list_flag)

                else:
                    inconnues_list_flag = 2
                    # print(f"Code {code} is not present in Thématiques_inconnues.")
                    Thématiques_inconnues_flags.append(inconnues_list_flag)

                if code in non_applicables_list:
                    non_applicables_list_flag = 1
                    # print(f"Code {code} is present in Thématiques_non_applicables.")
                    Thématiques_non_applicable_codes.append(code)
                    Thématiques_non_applicables_flags.append(non_applicables_list_flag)

                else:
                    non_applicables_list_flag = 2
                    # print(f"Code {code} is not present in Thématiques_non_applicables.")
                    Thématiques_non_applicables_flags.append(non_applicables_list_flag)

            Thématiques_inconnue_code_result = ','.join(Thématiques_inconnues_codes)
            logging.info("Thématiques_inconnues_codes_result--->", Thématiques_inconnue_code_result)
            # Step 1: Convert the string to a list
            codes_list = Thématiques_inconnue_code_result.split(',')

            # Step 2: Convert the list to a set to remove duplicates
            unique_codes_set = set(codes_list)

            # Step 3: Convert the set back to a string
            Thématiques_inconnues_codes_result = ','.join(unique_codes_set)
            logging.info("After removing duplicates:", Thématiques_inconnues_codes_result)

            Thématiques_non_applicable_code_result = ','.join(Thématiques_non_applicable_codes)
            logging.info("Thématiques_non_applicable_codes_result--->", Thématiques_non_applicable_code_result)
            # Step 1: Convert the string to a list
            codes_list = Thématiques_non_applicable_code_result.split(',')

            # Step 2: Convert the list to a set to remove duplicates
            unique_codes_set = set(codes_list)

            # Step 3: Convert the set back to a string
            Thématiques_non_applicable_codes_result = ','.join(unique_codes_set)
            logging.info("After removing duplicates:", Thématiques_non_applicable_codes_result)

            content, filtered_lines = interface_findlines(them_line, Thématiques_non_applicable_codes,
                                                          Thématiques_inconnues_codes,
                                                          Feps)
            logging.info("Thematics lines applicable to project or Arch from the DCI requirement------>", content, filtered_lines)
            # Create a list for the current iteration
            current_iteration_list = []

            # Append the values to the list
            current_iteration_list.append(dciReq)
            current_iteration_list.append(dciSignal)
            current_iteration_list.append(content)
            current_iteration_list.append('length of the thematics line applicable ' + str(filtered_lines))

            # Append the list to the master list
            master_list.append(current_iteration_list)
        except Exception as ex:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            exp_fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
            print(f"Thematics in the MUX sheet is not in the correct format.{ex} line no. {exc_tb.tb_lineno} file name: {exp_fname}")

    # Master_list e will have all the DCI reqs, signals, content and length
    print("interface master_list--------->", master_list)

    # Loop through each sublist in b
    for sublist in master_list:
        # Here if condition is used to check the thematic lines having zero means DCI req are not applicable
        if sublist[-1].endswith('length of the thematics line applicable 0'):
            # Use regular expression to extract the numeric part of sublist[-2]
            match = re.search(r'\d+', sublist[0])

            # If a match is found, append the formatted string to the new list c
            if match:
                # dcireq_not_applicable.append(f'{sublist[0].split("(")[0]}')
                dcireq_not_applicable.append(sublist[0])

        # Here if condition is used to check the thematic lines not equal to zero means DCI req are applicable
        # we are going to take the signal of each reqs for next process.
        if sublist[2].startswith('Please find the requirement signal for '):
            print("Req thematic is applicable.")
            signal = sublist[1]
            print("sublist[-1]------------>", signal)

            DCI_req = sublist[0]
            logging.info("DCI_req[-1]------------>", DCI_req)

            path = ICF.getInputFolder() + "\\" + EI.findInputFiles()[19]
            print("path ---->", path)
            Campagnec_Book = EI.openExcel(path)
            Campagnec_Book.activate()
            for sheet in Campagnec_Book.sheets:
                if sheet.name not in excluded_sheets:
                    if sheet.api.UsedRange.Find(signal):
                        print(f"Keyword '{signal}' found in sheet '{sheet.name}'")
                        if 'VSM' in sheet.name:
                            # Create a new list containing only elements containing 'VSM'
                            signal_present_sheetsss.append(sheet.name)
                            # Define the regex pattern for any four digits at the end
                            pattern = re.compile(r'_\d{4}$')
                            # Use list comprehension with regex condition
                            signal_present_sheet = sheet.name if pattern.search(sheet.name) else sheet.name.rsplit('_', 1)[0]
                            signal_present_sheetss.append(signal_present_sheet)

            print("signal_present_sheetsss------------>", signal_present_sheetsss)
            print("length of signal_present_sheetsss------------>", len(signal_present_sheetsss))

            print("signal_present_sheetss------------>", signal_present_sheetss)
            print("length of signal_present_sheetss------------>", len(signal_present_sheetss))

            signal_present_sheets = []

            for txt in signal_present_sheetss:
                logging.info("xtxttttttt-------->", txt)
                x = re.search(".{5}_.{2}_.{2}_.{2}_[^_]{4,5}", txt)
                logging.info("xxxxxxxxxxxxxxxxxxxxx----->", x)
                signal_present_sheets.append(x.group(0))

            signal_present_sheets = [*set(signal_present_sheets)]
            print("signal_present_sheets------------>", signal_present_sheets)
            print("length of signal_present_sheets a ------------>", len(signal_present_sheets))

            Campagne_sheet = Campagnec_Book.sheets['Synthèse des campagnes']
            sheet_value = Campagne_sheet.used_range.value
            N_Fiches_de_test_col = DCIPC.searchDataInExcelCache(sheet_value, 'N° Fiches de test')
            logging.info("N_Fiches_de_test_col---------->", N_Fiches_de_test_col)
            Synthèse_col = DCIPC.searchDataInExcelCache(sheet_value, 'Synthèse')
            logging.info("Synthèse_col---------->", Synthèse_col)

            try:
                # Create dictionaries to store columns for each row
                columns_N_Fiches = {pos[0]: pos[1] for pos in N_Fiches_de_test_col['cellPositions']}
                columns_Synthèse = {pos[0]: pos[1] for pos in Synthèse_col['cellPositions']}

                # Find the common row numbers
                common_row_numbers = set(columns_N_Fiches.keys()).intersection(columns_Synthèse.keys())

                # Get the matching cell positions along with their columns
                matching_positions = [(row, columns_N_Fiches[row], columns_Synthèse[row]) for row in common_row_numbers]

                # Extract colValue from the second tuple
                col_to_get_NT_NA_values = matching_positions[0][2]

                logging.info("col_to_get_NT_NA_values------->", col_to_get_NT_NA_values)
            except:
                print("Rows value of the N° Fiches de test and Synthèse are mismatched.")

            for sheet_name in signal_present_sheets:
                if sheet_name:
                    logging.info(f'Do something with sheet for signal: {sheet_name}')
                    Fun_name4 = EI.searchDataInColCache(sheet_value, 1, sheet_name)
                    logging.info("Fun_name4---------->", Fun_name4)
                    filtered_cell_values = [value for value in Fun_name4.get('cellValue', []) if value == sheet_name or value.startswith(sheet_name + '_')]

                    # Create a new dictionary with the same 'cellPositions' and the filtered 'cellValue'
                    filtered_dict = {'count': len(filtered_cell_values),
                                     'cellPositions': [pos for pos, cell_value in
                                                       zip(Fun_name4['cellPositions'], Fun_name4['cellValue'])
                                                       if cell_value in filtered_cell_values],
                                     'cellValue': filtered_cell_values}
                    # Print the result
                    logging.info("filtered_dict-->", filtered_dict)
                    if filtered_dict['cellPositions']:
                        for cellPositions in filtered_dict['cellPositions']:
                            row, col5 = cellPositions
                            logging.info("row, col5----------->", row, col5)
                            NT_NA_value_for_signal = EI.getDataFromCell(Campagne_sheet, (row, col_to_get_NT_NA_values))
                            print("NT_NA_value_for_signal----NT_NA_value_for_signal-->", NT_NA_value_for_signal)
                            if NT_NA_value_for_signal == 'NA':
                                NA_values_for_signal.append(NT_NA_value_for_signal)
                                NA_values_for_signal_sheet.append(sheet_name)
                                NA_values_for_req.append(DCI_req)
                                NA_DCI_signal.append(signal)
                            if NT_NA_value_for_signal == 'NT':
                                NT_values_for_signal.append(NT_NA_value_for_signal)
                                NT_values_for_signal_sheet.append(sheet_name)
                                NT_values_for_req.append(DCI_req)
                                NT_DCI_signal.append(signal)

            NA_signal_result = list(
                zip(NA_values_for_signal, NA_values_for_signal_sheet, NA_values_for_req, NA_DCI_signal))
            print("NA_signal_result------->", NA_signal_result)

            NT_signal_result = list(
                zip(NT_values_for_signal, NT_values_for_signal_sheet, NT_values_for_req, NT_DCI_signal))
            print("NT_signal_result------->", NT_signal_result)

            # Get the count of lists in each result
            na_count = len(NA_signal_result)
            nt_count = len(NT_signal_result)

            if (nt_count > 0 or na_count > 0 and nt_count > 0):
                logging.info(f"Data present only in NT_signal_result. NT_signal_result count: {nt_count}")

                # Example usage for the second element
                dcireqs_applicable = extract_unique_elements(NT_signal_result, 2)
                logging.info("dcireqs_applicable------->", dcireqs_applicable)

                # Example usage for the third element
                dcireqs_signal_applicable = extract_unique_elements(NT_signal_result, 3)
                logging.info("dcireqs_signal_applicable------->", dcireqs_signal_applicable)

                # if the req is present in the TP and req is not present in the DCI files
                if dci_req_not_found_in_file:
                    logging.info("hi")
                    first_row_in_table = Feps
                    second_row_in_table = "No"
                    third_row_in_table = 'For the ' + Feps + ' having the DCI Requirements ' + str(dci_req_results) + 'some DCI reqs ' + str(dci_req_not_found_in_file) + ' are missing from the DCI functional file.'
                    fourth_row_in_table = '--'
                    fifth_row_in_table = '--'
                    # Create a list and append the strings
                    rows_list = [first_row_in_table, second_row_in_table, third_row_in_table, fourth_row_in_table, fifth_row_in_table]
                    # Print the resulting list
                    logging.info("rows_list----------->", rows_list)
                    print("third_row_in_table- if condition--------->", third_row_in_table)

                # If the DCI req is present in PT and functional DCI files.
                # else:
                if dci_req_results == dci_req_results:
                    # if dci_req_results != None:
                    logging.info("NT_signal_result------->", NT_signal_result)
                    replace_sheet = ''
                    if NT_signal_result:
                        replace_sheet = NT_signal_result[0][1]
                    impact_sheet = impacted_sheet(dcireqs_applicable, Req_impacted_sheet)
                    third_row_in_table = 'For the ' + Feps + ' have the requirement ' + dcireqs_applicable + ' Requirement Associated with signal ' + dcireqs_signal_applicable + ' in the ' + Impacted_sheets + ' but this ' + Impacted_sheets + ' is NA for our Project ' + project_code + '. So this ' + dcireqs_applicable + ' Requirement is treatable in ' + replace_sheet + ' (Need to Raise QIA PT).'
                    print("third_row_in_table- else condition--------->", third_row_in_table)
                    # if 'having the thematic' in third_row_in_table and 'which is NA for our project' in third_row_in_table:
                    logging.info("havinghavinghavinghavinghavinghaving")
                    fourth_row_in_table = 'When QIA PT is treated this FEPS ' + Feps + ' can be treated.'
                    logging.info("modified_arch_for_table----------->", full_architecture)
                    modified_arch_for_table = full_architecture.replace(" ", "_")
                    fifth_row_in_table = "Can be treated"
                    second_row_in_table = "Yes"
                    first_row_in_table = Feps
                    # Create a list and append the strings
                    rows_list = [first_row_in_table, second_row_in_table, third_row_in_table, fourth_row_in_table, fifth_row_in_table]
                    # Print the resulting list
                    print("rows_list----------->", rows_list)
                    all_rows_list.append(rows_list)

            if na_count > 0 and nt_count == 0:
                logging.info("NT_signal_result------->", NA_signal_result)
                # Example usage for the second element
                dcireqs_applicable = extract_unique_elements(NA_signal_result, 2)
                print("dcireqs_applicable dfasd ------->", dcireqs_applicable)

                # Example usage for the third element
                dcireqs_signal_applicable = extract_unique_elements(NA_signal_result, 3)
                print("dcireqs_signal_applicable- asfsdfsd------>", dcireqs_signal_applicable)
                impact_sheet = impacted_sheet(dcireqs_applicable, Req_impacted_sheet)
                third_row_in_table = 'For the ' + Feps + ' have the requirement ' + dcireqs_applicable + ' Requirement Associated with signal ' + dcireqs_signal_applicable + ' in the ' + Impacted_sheets + ' but this ' + Impacted_sheets + ' which is NA for our Project ' + project_code + '. But there is no FT which are NT for our project to treat this DCI requirement. '
                print("third_row_in_table- else condition--------->", third_row_in_table)
                # if 'having the thematic' in third_row_in_table and 'which is NA for our project' in third_row_in_table:
                logging.info("havinghavinghavinghavinghavinghaving")
                fourth_row_in_table = 'When thematic are applicable for our project this FEPS can be treated.'
                logging.info("modified_arch_for_table----------->", full_architecture)
                modified_arch_for_table = full_architecture.replace(" ", "_")
                fifth_row_in_table = "'Cannot be tested because the thematics are not applicable for Arh"
                second_row_in_table = "NO"
                first_row_in_table = Feps
                # Create a list and append the strings
                rows_list = [first_row_in_table, second_row_in_table, third_row_in_table, fourth_row_in_table, fifth_row_in_table]
                # Print the resulting list
                print("rows_list----------->", rows_list)
                all_rows_list.append(rows_list)
            else:
                print("No data present in either list.")

            Campagnec_Book.save()

    # Join the list elements with commas and print the result
    dcireqs_not_applicable = ', '.join(dcireq_not_applicable)
    print("dcireqs_not_applicable result------------->", dcireqs_not_applicable)

    # DCI req is present in both and it is not applicable by the thematics.
    if dcireqs_not_applicable:
        third_row_in_table = 'For the ' + Feps + ' have the requirement ' + dcireqs_not_applicable + ' having the thematic ' + Thématiques_non_applicable_codes_result + ' which is NA for our project ' + project_code
        logging.info("third_row_in_table for dcireqs_not_applicable---------->", third_row_in_table)
        if 'having the thematic' in third_row_in_table and 'which is NA for our project' in third_row_in_table:
            logging.info("having")
            fourth_row_in_table = 'When thematic are applicable for our project this  ' + Feps + '  can be treated.'
            logging.info("modified_arch_for_table----------->", full_architecture)
            modified_arch_for_table = full_architecture.replace(" ", "_")
            fifth_row_in_table = "Cannot be tested because the thematics are not applicable for " + modified_arch_for_table
            second_row_in_table = "No"
            first_row_in_table = Feps
            # Create a list and append the strings
            rows_list = [first_row_in_table, second_row_in_table, third_row_in_table, fourth_row_in_table, fifth_row_in_table]
            # Print the resulting list
            print("rows_list---->", rows_list)
            all_rows_list.append(rows_list)

    if len(master_list) == 0:
        print("ooppooppo oppooppsdf")
        interface_requirements = ','.join(interface_requirements)
        third_row_in_table = 'For the ' + Feps + ' have all requirements ' + interface_requirements + ' are not present in Functional DCI or Mismatched with Arch.'
        fourth_row_in_table = '--'
        fifth_row_in_table = '--'
        second_row_in_table = "No"
        first_row_in_table = Feps
        # Create a list and append the strings
        rows_list = [first_row_in_table, second_row_in_table, third_row_in_table, fourth_row_in_table, fifth_row_in_table]
        # Print the resulting list
        print(rows_list)
        all_rows_list.append(rows_list)

    # Campagnec_Book.close()
    print("all__rows_list----------", all_rows_list)
    return all_rows_list


def NA_Table(Feps, Thématiques_inconnue, Thématiques_non_applicable, desired_architecture, project_code, full_architecture, impacetdSheets):
    logging.info("hi")
    table_contents = []
    path = ICF.getInputFolder() + "\\" + EI.findInputFiles()[1]
    logging.info("path ---->", path)
    test_Book = EI.openExcel(path)
    table_content = ''
    rows_list = ''
    if Feps:
        numeric_part = ''.join(filter(str.isdigit, Feps))
        logging.info("numeric_part--------->", numeric_part)
        NA_NT_values = []
        all_DCI_Reqs_sheets = []
        all_functional_Reqs_sheets = []
        try:
            test_Book.activate()
            Feps_sheet = test_Book.sheets['FEPS History']
            sheet_value = Feps_sheet.used_range.value
            Fun_name4 = EI.searchDataInColCache(sheet_value, 3, numeric_part.strip())
            row, col = Fun_name4['cellPositions'][0]
            time.sleep(2)
            logging.info(row, col)

            Impacted_sheets = EI.getDataFromCell(Feps_sheet, (row, col - 1))
            Impacted_sheets_reqs = EI.getDataFromCell(Feps_sheet, (row, col - 2))
            logging.info("Impacted_Requirements---------->", Impacted_sheets_reqs)
            logging.info("Impacted_Requirements.split()--------->", Impacted_sheets_reqs.split('\n'))
            Impacted_sheets_reqs = Impacted_sheets_reqs.split('\n')
            print(f'Impacted_Requirements for the FEPS_{numeric_part}---------->{Impacted_sheets_reqs}')
            # Specify the sheets to exclude from the output
            excluded_sheets = ['Matrice de tests', 'FEPS History', 'Impact', 'Historique Suppr', 'Feuil15']
            matching_requirements = []
            matching_requirements_sheets = []
            for req in Impacted_sheets_reqs:
                logging.info("req---->", req)
                # Iterate through all sheets and search for the keyword
                for sheet in test_Book.sheets:
                    if sheet.name not in excluded_sheets:
                        if sheet.api.UsedRange.Find(req):
                            print(f"Keyword '{req}' found in sheet '{sheet.name}'")
                            b = 'Keyword ' + req + ' found in sheet ' + sheet.name
                            if 'DCI' in req:
                                all_DCI_Reqs_sheets.append(b)
                            else:
                                all_functional_Reqs_sheets.append(b)

                            req_sheet = test_Book.sheets[sheet.name]
                            sheet_value06 = req_sheet.used_range.value
                            Search_req_in_sheet = DCIPC.searchDataInExcelCache(sheet_value06, req)
                            logging.info("Search_req_in_sheet------->", Search_req_in_sheet)
                            # Extract the cellValue from the search result
                            cell_value_list = Search_req_in_sheet.get('cellValue', [])

                            # Check if the list is not empty before accessing its first element
                            if cell_value_list:
                                cell_value = cell_value_list[0]

                                # Split the cellValue into a list of individual requirements
                                requirements_list = cell_value.split('|')

                                logging.info(f"Extracted requirements list: {requirements_list}")

                                # Check if req is in the list and print the matching requirement
                                matching_requirement = next(
                                    (requirement for requirement in requirements_list if req in requirement), None)

                                if matching_requirement:
                                    logging.info(f"Matching requirement: {matching_requirement}")
                                    matching_requirements.append(matching_requirement)
                                    matching_requirements_sheets.append(sheet.name)
                            else:
                                print("No cellValue found in the search result.")
                            break
                            # Clean up the requirements list
            cleaned_requirement = [requirement.strip().replace('\n', '') for requirement in matching_requirements]
            cleaned_requirements_impacted_sheet = [impacted_sheet.strip().replace('\n', '') for impacted_sheet in matching_requirements_sheets]

            # Define a regular expression pattern to match both formats
            dciRequirementPatterns = r'GEN-VHL-DCINT-[A-Za-z0-9_.]*\(\d\)|DCINT-\d{8}\(?\d{0,2}\)?'
            functionalReqPattern = r'REQ-\d{7}\s*(?:\(\w+\)|[A-Z]\b)|REQ_\w{4}_\w{3}_\w{3}_\w{3}\s*(?:\(\w+\)|[A-Z]\b).?|GEN-(?!.*(?:dci|DCI))'
            pattern = f"{dciRequirementPatterns}|{functionalReqPattern}"
            cleaned_requirements = [re.search(pattern, s).group() if re.search(pattern, s) else '' for s in cleaned_requirement]

            logging.info("Cleaned requirements list:", cleaned_requirements)
            logging.info(f'matching_requirements for the feps {Feps} requirements {cleaned_requirements}\n{cleaned_requirements_impacted_sheet}')
            Req_impacted_sheet = list(zip(cleaned_requirements, cleaned_requirements_impacted_sheet))
            logging.info("Req_impacted_sheet---------->", Req_impacted_sheet)
            # these logic is for of the reqs are not present in the TP

            # Extract the requirements without suffixes from cleaned_requirements
            cleaned_without_suffix = [req.split('(')[0] for req in cleaned_requirements]

            # Find the difference between Impacted_sheets_reqs and cleaned_without_suffix
            difference = list(set(Impacted_sheets_reqs) - set(cleaned_without_suffix))

            logging.info("Requirement is not present in the PT:::::", difference)

            # if reqs having DCI
            interface_requirements = []
            functional_requirements = []

            print("cleaned_requirements---------->", cleaned_requirements)
            # Remove empty elements from the list
            cleaned_requirements = [req for req in cleaned_requirements if req != '']

            # Now, cleaned_requirements will contain only non-empty elements
            print("cleaned_requirements new---------->",cleaned_requirements)


            for req in cleaned_requirements:
                if 'DCI' in req:
                    interface_requirements.append(req)
                else:
                    functional_requirements.append(req)
            print("all_DCI_Reqs_sheets----->", all_DCI_Reqs_sheets)
            print("TP feps sheet impacted sheets--------->", Impacted_sheets)

            print("Interface Requirements:", interface_requirements)
            print("Functional Requirements:", functional_requirements)

            if functional_requirements and functional_requirements is not None:
                rows_list = functionalreq(Feps, functional_requirements, Thématiques_inconnue,
                                          Thématiques_non_applicable, desired_architecture, project_code,
                                          full_architecture, Req_impacted_sheet, cleaned_requirements, Impacted_sheets)
            if interface_requirements:
                rows_list = interfacereq(Feps, interface_requirements, Thématiques_inconnue, Thématiques_non_applicable,
                                         excluded_sheets, desired_architecture, project_code, full_architecture,
                                         Req_impacted_sheet, all_DCI_Reqs_sheets, Impacted_sheets)

        except Exception as ex:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            exp_fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
            print(f"numeric_part is not present in the Testplan.{ex} line no. {exc_tb.tb_lineno} file name: {exp_fname}")
        table_contents.append(table_content)
    return rows_list


def updateFeps(projects, Fepss, Thématiques_inconnue, Thématiques_non_applicable, desired_architecture, project_code,
               full_architecture):
    inconnues_list_flag, non_applicables_list_flag = '', ''
    project = projects
    path = ICF.getInputFolder() + "\\" + EI.findInputFiles()[19]
    logging.info("path ---->", path)
    Campagnec_Book = EI.openExcel(path)
    path = ICF.getInputFolder() + "\\" + EI.findInputFiles()[1]
    logging.info("path ---->", path)
    test_Book = EI.openExcel(path)
    contents = []
    thematics_not_applicable_for_Feps = []

    global iteration_project
    global iteration_project_index
    for i,Feps in enumerate(Fepss):

        iteration_project_index = i

        numeric_part = ''.join(filter(str.isdigit, Feps))
        print("numeric_part--------->", numeric_part)
        NA_NT_values = []
        try:
            test_Book.activate()
            Feps_sheet = test_Book.sheets['FEPS History']
            sheet_value = Feps_sheet.used_range.value
            Fun_name4 = EI.searchDataInColCache(sheet_value, 3, numeric_part.strip())
            row, col = Fun_name4['cellPositions'][0]
            time.sleep(2)
            logging.info(row, col)
            Impacted_sheets = EI.getDataFromCell(Feps_sheet, (row, col - 1))
            if Impacted_sheets:
                Campagnec_Book.activate()
                Campagne_sheet = Campagnec_Book.sheets['Synthèse des campagnes']
                sheet_value = Campagne_sheet.used_range.value
                # Assuming Impacted_sheets is a list of sheet names
                logging.info("Impacted_sheets.split()--------->", Impacted_sheets.split('\n'))
                impacetdSheets = Impacted_sheets.split('\n')
                logging.info("impacetdSheetsimpacetdSheetsimpacetdSheetsimpacetdSheetsimpacetdSheets--------->", impacetdSheets)
                existing_values = []
                for sheet_name in impacetdSheets:
                    if sheet_name:
                        logging.info(f'Do something with sheet: {sheet_name}')
                        Fun_name4 = EI.searchDataInColCache(sheet_value, 1, sheet_name)
                        logging.info("Fun_name4---------->", Fun_name4)
                        # Extract only the cell values that exactly match the target string
                        filtered_cell_values = [value for value in Fun_name4.get('cellValue', []) if value == sheet_name or value.startswith(sheet_name + '_')]

                        # Create a new dictionary with the same 'cellPositions' and the filtered 'cellValue'
                        filtered_dict = {'count': len(filtered_cell_values),
                                         'cellPositions': [pos for pos, cell_value in
                                                           zip(Fun_name4['cellPositions'], Fun_name4['cellValue'])
                                                           if cell_value in filtered_cell_values],
                                         'cellValue': filtered_cell_values}

                        # Print the result
                        logging.info("filtered_dict-->", filtered_dict)

                        if filtered_dict['cellPositions']:
                            for cellPositions in filtered_dict['cellPositions']:
                                row, col5 = cellPositions
                                logging.info("row, col5----------->", row, col5)
                                feps = "ALTIS/ISUE vue(s) sur d'autres projets"
                                Feps_col = DCIPC.searchDataInExcelCache(sheet_value, feps)
                                logging.info("Feps_col----->", Feps_col)
                                if Feps_col['cellPositions']:
                                    row15, col25 = Feps_col['cellPositions'][0]
                                    logging.info("row, col15----------->", row15, col25)

                                    existing_value = EI.getDataFromCell(Campagne_sheet, (row, col25))
                                    logging.info("existing_value---------->", existing_value)

                                    # HERE I NEED TO IMPLEMENT THE FEPS LOGIC
                                    if existing_value:
                                        if existing_value is not None:
                                            if Feps not in existing_values:
                                                existing_value += "\n"  # Add a new line if there is an existing value
                                                new_value = existing_value + Feps
                                                # print("existing_value if condition-------->", new_value)
                                                lines = new_value.split("\n")
                                                unique_lines = set(lines)
                                                output_string = "\n".join(unique_lines)
                                                logging.info("output_string if condition- set ------->", output_string)
                                                EI.setDataFromCell(Campagne_sheet, (row, col25), output_string)
                                            else:
                                                print(f"FEPS value '{Feps}' is a repetition and not appended to cell.")
                                    else:
                                        EI.setDataFromCell(Campagne_sheet, (row, col25), Feps)
                                        logging.info("existing_value else condition-------->", existing_value)
                                logging.info("project0----------->", project)
                                project1 = DCIPC.searchDataInExcelCache(sheet_value, project)
                                # print("project1---------->", project1)
                                if project1['cellPositions']:
                                    row15, col15 = project1['cellPositions'][0]
                                    logging.info("row, col15----------->", row15, col15)
                                    logging.info("project1['cellValue'][1]----------->", project1['cellPositions'][0],  project1['cellValue'][0])
                                    if project == project1['cellValue'][0]:
                                        project_name = project1['cellValue'][0]
                                        logging.info(f'hiiii {project_name}')
                                        NA_NT_value = EI.getDataFromCell(Campagne_sheet, (row, col15))
                                        logging.info("NA_NT_values------->", NA_NT_value)
                                        NA_NT_values.append(NA_NT_value)
                        else:
                            print(f'Deleted the sheet impacted {sheet_name} by FEPS.')
                logging.info("NA_NT_values-->", NA_NT_values)
                count_NT = NA_NT_values.count('NT')
                count_NA = NA_NT_values.count('NA')
                content = 'For ' + Feps + ' will be ' + str(count_NT) + ' NT and ' + str(count_NA) + ' NA'
                contents.append(content)

                if NA_NT_values:
                    if all(value == 'NA' for value in NA_NT_values):
                        print("All elements are 'NA'")
                        row_list = NA_Table(Feps, Thématiques_inconnue, Thématiques_non_applicable, desired_architecture, project_code, full_architecture, impacetdSheets)
                        logging.info("result--------->", row_list)
                        thematics_not_applicable_for_Feps.append(row_list)
                    else:
                        print("It has a combination of 'NT' and 'NA' values")
            else:
                all_rows_list = []
                req_in_feps_sheet = EI.getDataFromCell(Feps_sheet, (row, col - 2))
                logging.info("req_in_feps_sheet------->", req_in_feps_sheet)
                # Split the string into individual requirements
                requirements_list = req_in_feps_sheet.split('\n')

                # Remove any newline characters from each requirement
                cleaned_requirements = [req.replace('\n', '') for req in requirements_list]

                # Join the cleaned requirements into a single string separated by commas
                feps_sheet_reqs = ','.join(cleaned_requirements)

                logging.info("feps_sheet_reqs-------->", feps_sheet_reqs)
                third_row_in_table = 'For the ' + Feps + ' have the requirement ' + feps_sheet_reqs + ', the requirement is deleted from the testplan.'
                logging.info("third_row_in_table- else --------->", third_row_in_table)
                fourth_row_in_table = 'The requirement is deleted from test plan.'
                fifth_row_in_table = "Can not be treated"
                second_row_in_table = "No"
                first_row_in_table = Feps
                # Create a list and append the strings
                rows_list = [first_row_in_table, second_row_in_table, third_row_in_table, fourth_row_in_table, fifth_row_in_table]
                # Print the resulting list
                print("rows_list----------->", rows_list)
                all_rows_list.append(rows_list)
                thematics_not_applicable_for_Feps.append(all_rows_list)

        except Exception as ex:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            exp_fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
            logging.info(f"{ex} line no. {exc_tb.tb_lineno} file name: {exp_fname}")
            print("Feps are not present in the testplan FEPS History sheet.")

    thematics_not_applicable_for_Feps = [elem for elem in thematics_not_applicable_for_Feps if elem]

    if thematics_not_applicable_for_Feps:
        print("table V------------thematics_not_applicable_for_Feps----->", thematics_not_applicable_for_Feps)
        print("entered")
        header = ["FEPS ID", "Can be Tested",
                  "If Yes, on which project\n\n\nIf No, list the thematic from requirement NA",
                  "Problem preventing FEPS to be tested on the requested projects", "Conclusion"]
        thematics_not_applicable_for_Feps.insert(0, header)
    else:
        print("NOt entered")

    project_text_bold = "\033[1m" + project + "\033[0m"
    print(project_text_bold)
    doc_content = [project + ':']
    NA_NT_content = doc_content + contents
    print("contentscontentscontents---------->", NA_NT_content)
    # global stop_threads
    # stop_threads = False
    # t1 = threading.Thread(target=excel_popup, args=(Campagnec_Book.name,))
    # t1.start()
    Campagnec_Book.save()
    stop_threads = True
    Campagnec_Book.close()
    test_Book.close()
    return NA_NT_content, inconnues_list_flag, non_applicables_list_flag, thematics_not_applicable_for_Feps


def getReqContent(Doc, req, ver):
    RqTable = ''
    # print(f"\nFinding the thematic..... {req,ver}")
    isDocValid = WDI.is_word_document_corrupted(Doc)
    if not isDocValid:
        TableList = WDI.getTables(Doc)
        if ((req.find('.') != -1) | (req.find('_') != -1)):
            req = req.replace('.', '-')
        RqTable = WDI.threading_findTable(TableList, req + "(" + ver + ")")
        print(f"RqTable123 {RqTable}")
        if RqTable != -1:
            chkOldFormat = WDI.checkFormat(RqTable, req + "(" + ver + ")")
            if chkOldFormat == 0:
                print(f"chkOldFormat111 {chkOldFormat}")
                Content = WDI.getOldContents(RqTable, req + "(" + ver + ")")
            else:
                Content = WDI.getReqContentData(RqTable, req + "(" + ver + ")")
        else:
            RqTable = WDI.threading_findTable(TableList, req + " (" + ver + ")")
            print(f"RqTable1234 {RqTable}")
            if RqTable != -1:
                chkOldFormat = WDI.checkFormat(RqTable, req + " (" + ver + ")")
                if chkOldFormat == 0:
                    print(f"chkOldFormat111 {chkOldFormat}")
                    Content = WDI.getOldContents(RqTable, req + " (" + ver + ")")
                else:
                    Content = WDI.getReqContentData(RqTable, req + " (" + ver + ")")
            else:
                RqTable = WDI.threading_findTable(TableList, req + " " + ver)
                print(f"RqTable1233333 {RqTable}")
                if RqTable != -1:
                    chkOldFormat = WDI.checkFormat(RqTable, req + " " + ver)
                    if chkOldFormat == 0:
                        print(f"chkOldFormat222 {chkOldFormat}")
                        Content = WDI.getOldContents(RqTable, req + " " + ver)
                    else:
                        Content = WDI.getReqContentData(RqTable, req + " " + ver)
                else:
                    RqTable = WDI.threading_findTable(TableList, req + "  " + ver)
                    if RqTable != -1:
                        chkOldFormat = WDI.checkFormat(RqTable, req + "  " + ver)
                        print(f"chkOldFormat3333 {chkOldFormat}")
                        if chkOldFormat == 0:
                            Content = WDI.getOldContents(RqTable, req + "  " + ver)
                        else:
                            Content = WDI.getReqContentData(RqTable, req + "  " + ver)
                    else:
                        Content = -1
    else:
        Content = -1
    return Content, RqTable


def findReqinC4Doc(req):
    searchResult = ''
    them_code = ''
    them_line = ''
    table = []
    Search_flag = 0
    reqName, reqVer = DCIPC.getReqVer(req)
    requirement_id = reqName + " " + reqVer
    print("in findC4req function")
    print("requirement_id--->", requirement_id)
    file_path = os.path.abspath(r'..\Output_Files\C4_Req_Content.docx')
    try:
        if os.path.exists(r'..\Output_Files\C4_Req_Content.docx'):
            logging.info("requirement_id-content13-------->", requirement_id)
            # content13 = WDI.getReqContentData(table, requirement_id)
            reqName, reqVer = DCIPC.getReqVer(requirement_id)
            # content13 = WDI.getReqContent(table, reqName, reqVer)
            content13, RqTable = getReqContent(file_path, reqName, reqVer)
            print("content13---->", content13)
            # raw_thematic = WDI.getRawThematic(RqTable, requirement_id)
            raw_thematic = WDI.getReqContentData(RqTable, requirement_id)
            print("raw_thematic----------->", raw_thematic)
            if RqTable == -1:
                return -1, -1
            else:
                logging.info("requirement_id-content13-------->", requirement_id)
                # # content13 = WDI.getReqContentData(table, requirement_id)
                # reqName, reqVer = DCIPC.getReqVer(requirement_id)
                # # content13 = WDI.getReqContent(table, reqName, reqVer)
                # content13, RqTable = getReqContent(file_path, reqName, reqVer)
                # print("content13---->", content13)
                # raw_thematic = WDI.getRawThematic(RqTable, requirement_id)
                # print("raw_thematic----------->", raw_thematic)
                # if RqTable != -1:
                if raw_thematic['effectivity']:
                    effective = str(raw_thematic['effectivity'].encode('utf-8').strip())
                    print(f'Effectivity is present for the req {requirement_id}')
                    # # print("Result effective= ", effective)
                    # # effectiveExpression = grepEffectiveExpression(effective)
                    # data = grepThematicsCode(effective)
                    thematics_code = BH.grepThematicsCode(effective)
                    print("thematics_code--------->", thematics_code)
                    them_line = BH.createCombination(thematics_code)
                    print("them_line---------->", them_line)
                    # Split the input string using the pipe character
                    them_code = them_line.split('|')
                    them_code = list(set([code for line in them_line.split('\n') for code in
                                          line.split('|')]))
                    print("them_code list:", them_code)
                if raw_thematic['diversity']:
                    print(f'Effectivity and lcdv is not present for the req {requirement_id} checking in diversity')
                    diversity = str(raw_thematic['diversity'].encode('utf-8').strip())
                    thematics_code = BH.grepThematicsCode(diversity)
                    logging.info("thematci_code--------->", thematics_code)
                    them_line = BH.createCombination(thematics_code)
                    logging.info("them_line---------->", them_line)
                    # Split the input string using the pipe character
                    # them_code = them_line.split('|')
                    them_code = list(
                        set([code for line in them_line.split('\n') for code in line.split('|')]))
                    logging.info("them_code list:", them_code)
                    print("req is present in the doc")
        else:
            return -1, -1
    except Exception as ex:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        exp_fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
        print(f"Error in getting the thematic part.{ex} line no. {exc_tb.tb_lineno} file name: {exp_fname}")
    return them_code, them_line


def findReqinSearchLogicDoc(req):
    them_code = ''
    them_line = ''
    reqName, reqVer = DCIPC.getReqVer(req)
    requirement_id = reqName + " " + reqVer
    print("requirement_id--->", requirement_id)
    file_path = os.path.abspath(r'..\Input_Files\Search_Logic')
    if os.path.exists(file_path) and os.path.isdir(file_path):
        # List all files in the folder
        files = os.listdir(file_path)

        # Filter for .docx files
        docx_files = [f for f in files if f.lower().endswith('.docx') or f.lower().endswith('.doc')]

        # Check if any .docx files are found
        if docx_files:
            for docx_file in docx_files:
                file_path = os.path.join(file_path, docx_file)
                logging.info("Found .docx file:", file_path)
                try:
                    if file_path:
                        logging.info("requirement_id-content13-------->", requirement_id)
                        reqName, reqVer = DCIPC.getReqVer(requirement_id)
                        content13, RqTable = getReqContent(file_path, reqName, reqVer)
                        print("content13---->", content13)
                        # raw_thematic = WDI.getRawThematic(RqTable, requirement_id)
                        raw_thematic = WDI.getReqContentData(RqTable, requirement_id)
                        print("raw_thematic----------->", raw_thematic)
                        if RqTable != -1:
                            if raw_thematic['effectivity']:
                                effective = str(raw_thematic['effectivity'].encode('utf-8').strip())
                                logging.info(f'Effectivity is present for the req {requirement_id}')
                                thematics_code = BH.grepThematicsCode(effective)
                                logging.info("thematics_code--------->", thematics_code)
                                them_line = BH.createCombination(thematics_code)
                                logging.info("them_line---------->", them_line)
                                # Split the input string using the pipe character
                                # them_code = them_line.split('|')
                                them_code = list(
                                    set([code for line in them_line.split('\n') for code in line.split('|')]))
                                logging.info("them_code list:", them_code)
                                print("req is present in the doc")
                                break
                                # elif raw_thematic['lcdv']:
                                #     print(f'Effectivity is not present for the req {requirement_id} checking in lcdv')
                                #     lcdv = str(raw_thematic['lcdv'].encode('utf-8').strip())
                                #     thematics_code = BH.grepThematicsCode(lcdv)
                                #     print("thematci_code--------->", thematics_code)
                                #     them_line = BH.createCombination(thematics_code)
                                #     print("them_line---------->", them_line)
                                #     # Split the input string using the pipe character
                                #     them_code = them_line.split('|')
                                #     print("them_code list:", them_code)
                            if raw_thematic['diversity']:
                                print(f'Effectivity and lcdv is not present for the req {requirement_id} checking in diversity')
                                diversity = str(raw_thematic['diversity'].encode('utf-8').strip())
                                thematics_code = BH.grepThematicsCode(diversity)
                                logging.info("thematci_code--------->", thematics_code)
                                them_line = BH.createCombination(thematics_code)
                                logging.info("them_line---------->", them_line)
                                # Split the input string using the pipe character
                                # them_code = them_line.split('|')
                                them_code = list(
                                    set([code for line in them_line.split('\n') for code in line.split('|')]))
                                logging.info("them_code list:", them_code)
                                print("req is present in the doc")
                                break
                            else:
                                print("Effectivity or diversity of req is not present in the doc")
                        else:
                            print("Requirement is not present in the search logic file.")
                            return -1, -1

                except Exception as ex:
                    exc_type, exc_obj, exc_tb = sys.exc_info()
                    exp_fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
                    # print(f"\nSearchLogic Output file not present in the folderrrrrrrr{ex} line no. {exc_tb.tb_lineno} file name: {exp_fname}")
                    print(f"\nSearchLogic Output file not present in the folder.")
                    return -1, -1
        else:
            print("No .docx or .doc files found in the folder.")
            return -1, -1
            # Add your logic here for when no .docx files are found
    else:
        print("Search_Logic Folder does not exist in Input Folder.")
        return -1, -1
    return them_code, them_line


def all_data(Task_ID, project, Fepss, ti, tna, nt, na, desired_architecture, full_architecture, desired_sheet_name, doc):
    All_contents = []
    # Iterate over the lists simultaneously using zip
    output_folder = ICF.getOutputFiles()
    # Create the full output path by joining the input folder with the document name
    output_path = os.path.join(output_folder, 'Output_document.docx')

    logging.info("project, Fepss, ti, tna, desired_architecture, full_architecture, desired_sheet_name------------->", project, Fepss, ti, tna, desired_architecture, full_architecture, desired_sheet_name)
    project_code = projectcode(desired_sheet_name, project)
    print("project_code-------->", project_code)
    no_of_NA_NT_contents, inconnues_list_flag, non_applicables_list_flag, thematics_not_applicable_for_Feps = updateFeps(project, Fepss, ti, tna, desired_architecture, project_code, full_architecture)
    result_tuples, Thématiques_inconnues_after_check = DCIPC.not_thematic(project, tna, ti, desired_architecture)
    # if Thématiques_inconnues_after_check != '':
    sheet_found_flag, file_name = to_find_sheet_inPLM(desired_sheet_name)
    NA_content = checkinplm(result_tuples, file_name, nt, na, desired_sheet_name, sheet_found_flag, Thématiques_inconnues_after_check, desired_architecture)
    logging.info("table after---------thematics_not_applicable_for_Feps----->", thematics_not_applicable_for_Feps)
    thematics_not_applicable_for_Feps = [elem for elem in thematics_not_applicable_for_Feps if elem]
    ALl_content = no_of_NA_NT_contents + NA_content + thematics_not_applicable_for_Feps
    All_contents.append(ALl_content)
    print("ALl_content-------->", ALl_content)
    All_feps_content = []
    for item in ALl_content:
        if isinstance(item, list):
            new_item = item[0] if item and isinstance(item[0], list) else item
            All_feps_content.append(new_item)
        else:
            All_feps_content.append(item)
    print("All_feps_content---------->", All_feps_content)
    append_to_word_document(All_feps_content, thematics_not_applicable_for_Feps, output_path, doc)
    fun(60)
    save_camp_output(Task_ID)


def save_camp_output(Task_ID):
    print("save_camp_outputsave_camp_output----------->")
    original_path = ICF.getInputFolder() + "\\" + EI.findInputFiles()[19]
    # Extracting directory and filename
    directory, filename = os.path.split(original_path)

    # Adding prefix
    new_filename = Task_ID+'_'+ filename
    print("new_filenamenew_filenamenew_filename----------->", new_filename)
    # Creating output folder path
    output_folder = ICF.getOutputFiles()

    # Making sure the output folder exists, if not, create it
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    # Constructing new path
    new_path = os.path.join(output_folder, new_filename)

    # You can now save the file to the new path
    print("New Path:", new_path)

    # Example of how to save a file (uncomment and adapt to your actual saving process)
    shutil.copy(original_path, new_path)


def summary_sheet():
    referentiel, referentiel_version = '', ''
    path = ICF.getInputFolder() + "\\" + EI.findInputFiles()[1]
    logging.info("path ---->", path)
    test_Book = EI.openExcel(path)
    test_Book.activate()
    Sommaire_sheet = test_Book.sheets['Sommaire']
    sheet_value = Sommaire_sheet.used_range.value
    Tp_reference = EI.searchDataInColCache(sheet_value, 1, 'Référence :')
    row, col = Tp_reference['cellPositions'][0]
    logging.info(row, col)
    tp_referentiel = EI.getDataFromCell(Sommaire_sheet, (row, col + 2))
    logging.info("tp_referentiel=========>", tp_referentiel)

    Fun_name13 = EI.searchDataInColCache(sheet_value, 1, 'Fonction :')
    row, col = Fun_name13['cellPositions'][0]
    logging.info(row, col)
    Domaine_name = EI.getDataFromCell(Sommaire_sheet, (row - 1, col + 2))
    Function_name = EI.getDataFromCell(Sommaire_sheet, (row, col + 2))
    PT_Version = EI.getDataFromCell(Sommaire_sheet, (row + 2, col))
    logging.info("Domaine_name=========>", Domaine_name)
    logging.info("Fun_name=========>", Function_name)
    logging.info("PT_Version=========>", PT_Version)
    PT_Version = f"V{int(PT_Version)}"
    logging.info("after change PT_Version=========>", PT_Version)

    Function_Name = f"{Domaine_name.split(' ')[0]}_{Function_name.split(' ')[0]} - {Function_name.split(' ')[-1]}"
    logging.info("Function_Name=========>", Function_Name)
    # Split the string at the hyphen and get the part after it
    Func_name = Function_name.split('-')[1].strip()
    logging.info("Func_name result-------->", Func_name)

    referentiel_keywords = ['Referentiel', 'Réferentiel', 'Référentiel']

    for keyword in referentiel_keywords:
        try:
            referentiel_name = EI.searchDataInColCache(sheet_value, 8, keyword)
            row, col = referentiel_name['cellPositions'][0]
            logging.info(row, col)
            referentiel_version = EI.getDataFromCell(Sommaire_sheet, (row + 1, col + 1))

            referentiel = EI.getDataFromCell(Sommaire_sheet, (row + 1, col + 2))
            logging.info("referentiel_name---------->", referentiel, referentiel_version)
            break  # Break out of the loop if found
        except Exception as ex:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            exp_fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
            print(f"\nSearchLogic Output file not present in the folderrrrrrrr{ex} line no. {exc_tb.tb_lineno} file name: {exp_fname}")
        #
        # except:
            continue  # Continue to the next keyword if not found
    return tp_referentiel, PT_Version, Function_Name, Func_name, referentiel, referentiel_version


def campagne_Ver_main(name, start_date, Task_ID, reference, projects, project_details):
    print("Tool Started")
    global iteration_project
    start = time.time()
    # Name = input("Enter the Name : ")
    # Start_date = input("Enter the Start_date : ")
    # Fepss = input("Enter the GAELE_reference by giving , : ")
    # Fepss = Fepss.split(',')
    # # Prompt the user to enter the number of projects
    # num_projects = int(input("Enter the number of projects: "))
    # Initialize empty lists to store the desired sheets and architectures
    Name = name
    Start_date = start_date
    task_ID = Task_ID
    Fepss = reference
    Fepss = Fepss.split(',')
    # to_address_mail = to_address
    # to_addres = to_address_mail.split(',')
    # Prompt the user to enter the number of projects
    num_projects = projects
    # Initialize empty lists to store the desired sheets and architectures

    desired_sheets = []
    desired_architectures = []
    desired_full_architectures = []
    fun(0)
    print("project_details-------->", project_details)

    # Iterate over project details
    for project in project_details:
        sheet_name = project['sheet_name']
        architecture = project['architecture']

        # Append to respective lists
        desired_sheets.append(sheet_name)
        desired_architectures.append(architecture)
        desired_full_architectures.append(architecture)

    # # Loop to get input for each project
    # for i in range(num_projects):
    #     sheet_name = input(f"Enter the desired sheet name for project {i + 1}: ")
    #     architecture = input(f"Enter the desired architecture for project {i + 1}: ")
    #     desired_full_architectures.append(architecture)
    #     desired_sheets.append(sheet_name)
    #     desired_architectures.append(architecture)

    # Print the list of desired sheets and architectures
    print("Desired sheets:", desired_sheets)
    print("Desired architectures:", desired_architectures)
    print("Desired full_architectures:", desired_full_architectures)
    fun(5)
    thematics_not_applicable_for_Feps = ''
    projects, Thématiques_inconnues, Thématiques_non_applicables, NT_values, NA_values = get_values_from_campagne(Name, Start_date)
    logging.info("projects, Thématiques_inconnues, Thématiques_non_applicables, NT_values, NA_values--------->",
                 projects, Thématiques_inconnues, Thématiques_non_applicables, NT_values, NA_values)
    fun(15)
    # Iterate over the lists simultaneously using zip
    output_folder = ICF.getOutputFiles()
    # Create the full output path by joining the input folder with the document name
    output_path = os.path.join(output_folder, 'Output_document.docx')
    doc = Document()
    result = ""
    tp_referentiel, PT_Version, Function_Name, TP_Func_name, referentiel, referentiel_version = summary_sheet()
    global iteration_project
    iteration_project = [True]*len(Fepss)
    if len(projects) > 1 and num_projects == 1:
        print("Entered only one project if the Campagne having 2 projects.")
        fun(20)
        for project in projects:
            sheet_name_parts = sheet_name.split('_')
            project_parts = project.split('_')
            logging.info("sheet_name_parts--------->", sheet_name_parts)
            common_parts = set(sheet_name_parts) & set(project_parts)
            logging.info("common_parts---------->", common_parts)
            # If there is a common part, extract it
            if common_parts:
                logging.info("project----->", project)
                # Find the index of the project
                index = projects.index(project)
                logging.info("indexindexindex---------->", index)
                # Use the index to access the corresponding elements from other lists
                Thématiques_inconnues_value = Thématiques_inconnues[index]
                Thématiques_non_applicables_value = Thématiques_non_applicables[index]
                NT_values_value = NT_values[index]
                NA_values_value = NA_values[index]
                logging.info(
                    "Thématiques_inconnues_value, Thématiques_non_applicables_value, NT_values_value, NA_values_value, desired_sheet_name, desired_architecture-------> ",
                    Thématiques_inconnues_value, Thématiques_non_applicables_value, NT_values_value, NA_values_value,
                    desired_sheets[0], desired_architectures[0])

                if project == project:
                    logging.info("desired_sheet_name,desired_architecture, project_name-------->", project,
                                 desired_sheets[0], desired_architectures[0])
                    nt = NT_values_value
                    na = NA_values_value
                    desired_architecture = desired_architectures[0]
                    full_architecture = desired_full_architectures[0]
                    desired_sheet_name = desired_sheets[0]
                    fun(30)
                    logging.info("project, Fepss, Thématiques_inconnues_value, Thématiques_non_applicables_value, nt=str(NT_values_value[0]), na=str(NA_values_value[0]), desired_architecture=desired_architectures[0], full_architecture=desired_full_architectures[0], desired_sheet_name=desired_sheets[0]0------------->",
                        project, Fepss, Thématiques_inconnues_value, Thématiques_non_applicables_value, nt, na,
                        desired_architecture, full_architecture, desired_sheet_name)
                    all_data(Task_ID, project, Fepss, Thématiques_inconnues_value, Thématiques_non_applicables_value, nt, na, desired_architectures[0], desired_full_architectures[0], desired_sheets[0], doc)
    # fun(40)
    if len(projects) > 1 and num_projects > 1:
        print("Entered more than one Project")
        fun(20)
        project_data = zip(projects, Thématiques_inconnues, Thématiques_non_applicables, NT_values, NA_values,
                           desired_sheets, desired_architectures, desired_full_architectures)
        for project, ti, tna, nt, na, desired_sheet_name, desired_architecture, full_architecture in project_data:
            logging.info("desired_sheet_name,desired_architecture, project_name-------->", desired_sheet_name,
                         desired_architecture, project)
            logging.info(
                "project, ti, tna, nt, na, desired_sheet_name, desired_architecture, full_architecture-------->",
                project, ti, tna, nt, na, desired_sheet_name, desired_architecture, full_architecture)
            all_data(Task_ID, project, Fepss, ti, tna, nt, na, desired_architecture, full_architecture, desired_sheet_name, doc)

    if len(projects) == 1 and num_projects == 1:
        print("Entered one Project and Campagne having one project")
        fun(20)
        project_data = zip(projects, Thématiques_inconnues, Thématiques_non_applicables, NT_values, NA_values,
                           desired_sheets, desired_architectures, desired_full_architectures)
        for project, ti, tna, nt, na, desired_sheet_name, desired_architecture, full_architecture in project_data:
            logging.info("desired_sheet_name,desired_architecture, project_name-------->", desired_sheet_name,
                         desired_architecture, project)
            logging.info(
                "project, ti, tna, nt, na, desired_sheet_name, desired_architecture, full_architecture-------->",
                project, ti, tna, nt, na, desired_sheet_name, desired_architecture, full_architecture)
            all_data(Task_ID, project, Fepss, ti, tna, nt, na, desired_architecture, full_architecture, desired_sheet_name, doc)
    fun(90)
    print(f"Word document saved to: {output_path}")
    fun(100)
    window = pgw.getWindowsWithTitle("VSM_PT")
    # window[0].maximize()
    # window[0].activate()
    window[0].activate()
    window[0].restore()
    end1 = time.time()
    print("\nexecution time " + str(end1 - start))
    print("Task Fully Completed")


if __name__ == '__main__':
    ICF.loadConfig()
    campagne_Ver_main()
