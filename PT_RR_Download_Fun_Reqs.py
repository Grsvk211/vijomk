import datetime
import shutil
import docx
import sys
import ExcelInterface as EI
import os
import re
import InputConfigParser as ICF
import WordDocInterface as WDI
import web_interface as WI
import NewRequirementHandler as NRH
import threading
import pygetwindow as pgw
import time
import KeyboardMouseSimulator as KMS
import logging
from itertools import zip_longest
from DCI_download_webinterface import startDocumentDownload
from itertools import chain
from docx import Document
from docx.enum.text import WD_BREAK
import supporting_requirement as SR
date_time = datetime.datetime.now()


def validateInputDocument(ipDocName):
    refNum = re.findall(
        "([A-Z0-9]{4,5})+(_[0-9]{2})+(_[A-Z0-9]{4,5})+|([A-Z0-9]{4})+(_[A-Z0-9]{4})+(_[A-Z0-9]{3})+(_[0-9]{3})",
        ipDocName)
    verNum = re.findall("[vV]{1}[0-9]{1,2}\.[0-9]{1,2}", ipDocName)
    if refNum:
        refNum = list(set(refNum))
        logging.info(refNum)
        verNum = list(set(verNum))
        logging.info("verNum", verNum)
        logging.info("Version Count =", len(verNum))
        logging.info("Ref Num Count =", len(refNum))
        status = False
        if len(refNum) == 1 and len(verNum) == 1:
            status = True
        return status


def addDataInDocument(data, type):
    logging.info(f"\n-------> data <-------\n{data}")
    output_dir = os.path.abspath(r"..\Output_Files")
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    dd = os.path.exists(r'..\Output_Files\C4_Req_Content.docx')
    if os.path.exists(r'..\Output_Files\C4_Req_Content.docx'):
        document = Document(r'..\Output_Files\C4_Req_Content.docx')
    else:
        document = Document()

    is_empty = is_document_empty(document)
    if not is_empty and type == 'H':
        run = document.add_paragraph().add_run()
        run.add_break(WD_BREAK.PAGE)

    # Add content to the document
    if type == 'H':
        document.add_heading(f"--------------- {data} -------------", level=1)
    elif type == 'T':
        document.add_paragraph('\n')
        table = data[0]  # Extracting the first element from the tuple
        document.add_paragraph(f"File Name: {data[1]}\n")

        destination_table = document.add_table(rows=len(table.rows), cols=len(table.columns))
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                destination_table.cell(row_idx, col_idx).text = cell.text
    else:
        document.add_paragraph(data)

    output_dir = os.path.abspath(r"..\Output_Files")
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    savingPath = os.path.abspath(r'..\Output_Files\C4_Req_Content.docx')
    # Save the document
    document.save(savingPath)
    return 1


def is_document_empty(doc):
    # Check if the document contains any paragraphs, tables, or headings
    has_paragraphs = bool(doc.paragraphs)
    has_tables = bool(doc.tables)
    has_headings = any(p.style.name.startswith('Heading') for p in doc.paragraphs)

    # The document is considered empty if there are no paragraphs, tables, or headings
    return not (has_paragraphs or has_tables or has_headings)


def find_and_save_keyword(file_path, keyword, Keyword_file_folder):
    try:
        # Open the Word document
        doc = Document(file_path)

        # Check if the keyword is present in the document
        if any(keyword in paragraph.text for paragraph in doc.paragraphs):
            # Create the output folder if it doesn't exist
            if not os.path.exists(Keyword_file_folder):
                os.makedirs(Keyword_file_folder)

            # Extract the filename from the original path
            file_name = os.path.basename(file_path)

            # Create the output file path
            output_path = os.path.join(Keyword_file_folder, file_name)

            # Save the document to the output folder
            doc.save(output_path)

            print(f"Keyword found in '{file_name}'. Document saved to '{Keyword_file_folder}'.")
    except Exception as e:
        print(f"Error processing '{file_path}': {e}")


def process_folder(input_folder, keyword, Keyword_file_folder):
    # Get a list of all files in the input folder
    files = [f for f in os.listdir(input_folder) if f.endswith(".docx")]

    # Process each file in the input folder
    for file in files:
        file_path = os.path.join(input_folder, file)
        find_and_save_keyword(file_path, keyword, Keyword_file_folder)


def getTables(fileName):
    wordDoc = docx.Document(fileName)
    # print("wordDoc.tables - ", wordDoc.tables)
    return wordDoc.tables


def find_tables_by_keyword_and_header(table, keyword, header):
    matching_tables = []

    # for table in doc.tables:
    # for table in table_lists:
    # Check if the first cell in the first row contains the specified header
    if table.cell(0, 0).text.strip().lower() == header.lower():
        for row in table.rows:
            for cell in row.cells:
                if keyword in cell.text:
                    # print("kokokokop")
                    matching_tables.append(table)
                    break
            if matching_tables:
                break
    return matching_tables


def getCellAbsVal(sheet, row, col):
    for i in range(row, 0, -1):
        cellVal = EI.getDataFromCell(sheet, f"{col}{i}")
        if cellVal is not None:
            return cellVal
    return None


def get_values_until_vpt(sheet, last_row):
    values_list = []
    for row in range(last_row, 0, -1):
        cell_value = sheet.range(f"B{row}").value
        if cell_value == "VPT":
            break
        values_list.append(cell_value)
    return values_list


def get_values_before_vpt(sheet, last_row):
    values_list = []
    # last_row = sheet.cells.last_cell.row
    for row in range(last_row, 0, -1):
        cell_value = sheet.range(f"B{row}").value
        if cell_value == "V PT":
            break
        values_list.append(cell_value)
    # Remove None values and reverse the list
    values_list = [value for value in reversed(values_list) if value is not None]
    return values_list


def getReqVer(req):
    if req.find('(') != -1:
        new_reqName = req.split("(")[0].split()[0] if len(req.split("(")) > 0 else ""
        new_reqVer = req.split("(")[1].split(")")[0] if len(req.split("(")) > 1 else ""
    else:
        new_reqName = req.split()[0] if len(req.split()) > 0 else ""
        new_reqVer = req.split()[1] if len(req.split()) > 1 else ""
    return new_reqName.strip(), new_reqVer.strip()


def parseIpDocId_ver(docvalues):
    # print("parseIpDocId_ver +")
    feps = docvalues
    logging.info("Feps Doc = ", feps)
    docs = []
    for fep in feps:
        if fep:
            logging.info(fep)
            tup = ()

            x = re.search("([A-Z0-9]{4,5})(_[0-9]{2})(_[A-Z0-9]{4,5})", fep)
            if x is not None:
                x = x.group()
                pattern = r'([_|\s][vV]{1}[0-9]{1,2}.[0-9]{0,2})|([_|\s][vV]{1}[0-9]{1,2})'
                y = re.split(pattern, fep)
                logging.info("y = ", y)
                if len(y) > 1:
                    y[1] = y[1].replace("_", "")
                    y[1] = y[1].strip()
                    logging.info("Y[1] = ", y[1])

                    tup = tup + (x, y[1])
                    docs.append(tuple(tup))
                    del tup
                else:
                    print("No Version found in input document name")
    # print("parseIpDocId_ver -")
    return docs


def parseIpDocName(docvalues):
    feps = docvalues
    docs = []
    for fep in feps:
        if fep:
            # if not validateInputDocument(fep):
            #     return -1
            inner = []
            # for f in fep:
            #  print(fep)
            x = fep.split(" - ")
            inner.append(x[0])
            docs.append(inner)
            del inner
    # print("+++++++++++++++++docs=" + str(docs))
    return docs


def searchAnalyseDeEntrant(b, reqFnd):
    # Open Excel
    final_result = []
    Analyse_de_entrant = EI.findInputFiles()[2]
    logging.info("analyse = ", Analyse_de_entrant, len(Analyse_de_entrant))
    alldocs = []
    all_docs_all_docs= []
    if len(Analyse_de_entrant) != 0:
        for i, analyseSheet in enumerate(Analyse_de_entrant):
            try:
                if not analyseSheet.startswith("BSI"):  # Check if filename starts with "BSI"
                    logging.info("Sheets=========>", analyseSheet)
                    global stop_threads
                    stop_threads = False
                    t1 = threading.Thread(target=excel_popup, args=(analyseSheet,))
                    t1.start()
                    analyseDeEntrant = EI.openExcel(ICF.getInputFolder() + "\\" + analyseSheet)
                    stop_threads = True
                    logging.info("analyse = ", analyseDeEntrant.sheets)
                    for sheet in analyseDeEntrant.sheets:
                        logging.info("Sheet name = ", sheet.name)
                        sheet_value = sheet.used_range.value
                        for requirement_id in b:
                            logging.info("requirement_id---->", requirement_id)
                            reqName, reqVer = NRH.getReqVer(requirement_id)
                            logging.info("Task Name = ", reqName)
                            result = EI.searchDataInExcelCache(sheet_value, reqName)
                            if result['count'] > 0:
                                logging.info("result---------->", result)
                                search_formats = get_search_formats(reqName, reqVer)
                                for search in search_formats:
                                    logging.info("search--------->",  search)
                                    logging.info("search1--------->",  result['cellValue'][0])
                                    for symbol in ['->', '=>', '-->', '==>']:
                                        if symbol in result['cellValue'][0]:
                                            logging.info(f"Arrow '{symbol}' is present.")
                                            split_string = result['cellValue'][0].split(symbol)  # Split the string based on the arrow symbol
                                            if len(split_string) > 1:
                                                result['cellValue'][0] = split_string[1].strip()
                                                if search == result['cellValue'][0]:
                                                    logging.info("if search == result['cellValue'][0]", search, result['cellValue'][0])
                                                    logging.info("ooooooooooooooooooooooooooo")
                                                    colno = result['cellPositions'][0][1]
                                                    docValues = sheet.range((10, colno), (12, colno)).value
                                                    docIdVer = parseIpDocId_ver(docValues)
                                                    docNames = parseIpDocName(docValues)
                                                    final_result = list(zip_longest(docIdVer, docNames))
                                                    logging.info("final result if", final_result)
                                                    alldocs.append(final_result)
                                                    # Adding the requirement ID to the tuple
                                                    final_result = [(tup[0], tup[1], requirement_id) for tup in
                                                                    final_result]
                                                    all_docs_all_docs.append(final_result)
                                                    alldocs.append(final_result)
                                                    reqFnd.append(requirement_id)
                                                    # break
                                        else:
                                            if search == result['cellValue'][0]:
                                                logging.info("else search == result['cellValue'][0]", search, result['cellValue'][0])
                                                logging.info("ooooooooooooooooooooooooooo")
                                                colno = result['cellPositions'][0][1]
                                                docValues = sheet.range((10, colno), (12, colno)).value
                                                docIdVer = parseIpDocId_ver(docValues)
                                                docNames = parseIpDocName(docValues)
                                                final_result = list(zip_longest(docIdVer, docNames))
                                                logging.info("final result else", final_result)
                                                alldocs.append(final_result)
                                                # Adding the requirement ID to the tuple
                                                final_result = [(tup[0], tup[1], requirement_id) for tup in final_result]
                                                all_docs_all_docs.append(final_result)
                                                reqFnd.append(requirement_id)
                                    logging.info("all_docs_all_docs----------->", all_docs_all_docs)
                                                # break
                    analyseDeEntrant.close()
            except Exception as ex:
                exc_type, exc_obj, exc_tb = sys.exc_info()
                exp_fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
                print(f"\nSomething  wrong {ex} line no. {exc_tb.tb_lineno} file name: {exp_fname}")
        if alldocs:
            return alldocs
        print("Input keyword not found in Analyse de entrant")
        return [-1]
    else:
        print("No analyse de entrant")
        return [-1]


def getContent(Doc, ReqName, ReqVer):
    logging.info("Doc,ReqName,ReqVe------->", Doc, ReqName, ReqVer)
    Content = ''
    RqTable = ''
    try:
        TableList = getTables(Doc)
        logging.info("TableList------>", TableList)
        # RqTable=threading_findTable(TableList, ReqName+"("+ReqVer+")")
        RqTable = WDI.threading_findTable(TableList, ReqName)
        logging.info("")
        if RqTable == -1:
            if ((ReqName.find('.') != -1) | (ReqName.find('_') != -1)):
                ReqName = ReqName.replace('.', '-')
            RqTable = WDI.threading_findTable(TableList, ReqName + "(" + ReqVer + ")")
        else:
            RqTable = WDI.threading_findTable(TableList, ReqName + "(" + ReqVer + ")")
        if RqTable != -1:
            chkOldFormat = WDI.checkFormat(RqTable, ReqName + "(" + ReqVer + ")")
            if chkOldFormat == 0:
                Content = WDI.getOldContents(RqTable, ReqName + "(" + ReqVer + ")")
            else:
                Content = WDI.getNewContents(RqTable, ReqName + "(" + ReqVer + ")")
        else:
            RqTable = WDI.threading_findTable(TableList, ReqName + " " + ReqVer)
            if RqTable != -1:
                chkOldFormat = WDI.checkFormat(RqTable, ReqName + " " + ReqVer)
                if chkOldFormat == 0:
                    Content = WDI.getOldContents(RqTable, ReqName + " " + ReqVer)
                else:
                    Content = WDI.getNewContents(RqTable, ReqName + " " + ReqVer)
            else:
                RqTable = WDI.threading_findTable(TableList, ReqName + "  " + ReqVer)
                if RqTable != -1:
                    chkOldFormat = WDI.checkFormat(RqTable, ReqName + "  " + ReqVer)
                    if chkOldFormat == 0:
                        Content = WDI.getOldContents(RqTable, ReqName + "  " + ReqVer)
                    else:
                        Content = WDI.getNewContents(RqTable, ReqName + "  " + ReqVer)
                else:
                    Content = -1
        logging.info("content123--->", RqTable, Content, Doc)
        req_content_tuples = list(zip(RqTable, Content, Doc))
    except Exception as ex:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        exp_fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
        print(f"\nSomething went wrong {ex} line no. {exc_tb.tb_lineno} file name: {exp_fname}")
    return RqTable, Content, Doc


def get_search_formats(reqName, reqVer):
    formats = [
        reqName + " " + reqVer,
        reqName + "  " + reqVer,
        reqName + " (" + reqVer + ")",
        reqName + "(" + reqVer + ")",
    ]
    return formats


def getdocs(b, reqFound):
    if os.path.exists(r'..\Input_Files\Analyse_des_entrant_Global_suite.xlsx'):
        print("Analyse_des_entrant present in the Input Folder.")
        pass
    else:
        print("Analyse_des_entrant not present in the Input Folder. Downloading the Analyse_des_entrant from Doc Info.")
        startDocumentDownload([("01272_19_02283", "")], True, False, ('xlsx',))
    all_docs = []
    logging.info("In getdocs function.")
    try:
        print("Searching the Requirements in the Global and VSM Analyse_des_entrant.")
        all_docs = searchAnalyseDeEntrant(b, reqFound)
        logging.info("all_docsall_docsall_docs----------------->", all_docs)

        if all_docs != -1:
            if -1 in all_docs:
                logging.info("Proceed with another condition")
            else:
                logging.info("All conditions checked, no -1 found")
            logging.info("all_docs----->", all_docs)

            # # Remove -1 and empty lists from all_docs
            all_docs = list(filter(lambda x: x != -1 and x != [], all_docs))
            logging.info("after removing -1 in the all_docs list----->", all_docs)

            valid_docs = [result for result in all_docs if result[0][0] is not None]
            none_docs = [result[0][1][0] for result in all_docs if result[0][0] is None]
            logging.info("Valid Documents:", valid_docs)
            # all_docs = [[(('02017_19_02197', 'V7.0'), ['SSD_HMIF_GROUND_LINK_HMI']), (('02017_19_02197', 'V7.0'), ['ALLOC_MATRIX_SSD_GDLK'])], [(('01991_19_01457', 'V13.0'), ['EEAD_DCT_CABIN_DETECT_CRASH_TRIGGER22Q4_01991_19_01457 V13.0 (IR DCROSS (PSE))'])], [(('01843_21_00859', 'V1.0'), ['NT_HMIF_VSM_BSI_PCGA_ISS-0119182_Altis-10816290_Dates_management_for_StartEnd_of_Trip'])]]

            flat_valid_docs = list(chain.from_iterable(valid_docs))
            logging.info("Flattened Valid Documents:", flat_valid_docs)

            # Flatten the nested list
            flat_list = []

            for sublist in flat_valid_docs:
                for item in sublist:
                    try:
                        if item and len(item) >= 2 and item[0] is not None:
                            # print("Debug - item[0]:", item[0])  # Print for debugging
                            # print("Debug - item[1]:", item[1])  # Print for debugging

                            if isinstance(item[1], list) and item[1]:
                                flattened_item = (item[0], item[1][0].lstrip('v').lstrip('V'))
                            elif isinstance(item[1], str):
                                flattened_item = (item[0], item[1].lstrip('v').lstrip('V'))
                            else:
                                # print("Debug - item[1] is neither list nor string:", item[1])
                                continue  # Skip processing for unexpected item[1] type

                            flat_list.append(flattened_item)
                    except IndexError as e:
                        logging.info("Error - IndexError:", e)
                        logging.info("Error - item:", item)

            # Print intermediate result
            logging.info("Intermediate flat_list:", flat_list)

            # Remove duplicates
            unique_list = list(set(flat_list))

            # Output the result
            result_list = [(item[0], item[1]) for item in unique_list]
            logging.info("result_list------>", result_list)

            logging.info("None Documents:", none_docs)
            result_list_none = []
            pattern = re.compile(r'(\d{5}_\d{2}_\d{5})-[vV](\d+(?:\.\d+)?)', re.IGNORECASE)
            unique_documents = list(set(none_docs))
            # # Example usage
            # text = 'SSVS_SSFD_GEN2_RSP_SRM_SUPERVISE_RECONFIGURATION_AND_MAINTENANCE_22Q3_02014_19_00638-v12(LA-00417277)'
            for text in unique_documents:
                match = pattern.search(text)
                if match:
                    version_info = (match.group(1), f'V{float(match.group(2)):.1f}')
                    logging.info(version_info)
                    result_list_none.append(version_info)
            logging.info("result_list_none version-------->", result_list_none)

            # Combine the two lists
            combined_list = result_list_none + result_list

            # Use a dictionary to keep track of unique keys and versions
            unique_dict = {}
            for item in combined_list:
                key, version = item
                if key not in unique_dict:
                    unique_dict[key] = {version}
                else:
                    unique_dict[key].add(version)
            # Convert the dictionary values back to a list
            result_combined = [(key, v) for key, versions in unique_dict.items() for v in versions]
            logging.info("Downloading the Requirement input Documents which are found in the Analyse_des_entrant.")
            logging.info("result_combined--------->", result_combined)

            result_combined_modified = [(x[0][:14], x[1]) for x in result_combined]

            logging.info(result_combined_modified)

            input_folder = ICF.getInputFolder()
            doc_files_present = any(file.endswith('.doc') or file.endswith('.docx') for file in os.listdir(input_folder))
            if not doc_files_present:
                # for i in result_combined:
                for i in result_combined_modified:
                    WI.startDocumentDownload([i])

    except Exception as ex:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        exp_fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
        print(f"\nSomething went wrong {ex} line no. {exc_tb.tb_lineno} file name: {exp_fname}")
    return all_docs


def organize_word_docs(folder_path):
    # List all files in the folder
    files = os.listdir(folder_path)

    # Check if a Word document is present in the folder
    word_docs = [file for file in files if file.endswith('.doc') or file.endswith('.docx') or file.endswith('.rtf')]

    if word_docs:
        # Create a new folder for Word documents
        word_folder_path = os.path.join(folder_path, 'Word_Documents')
        os.makedirs(word_folder_path, exist_ok=True)

        # Move Word documents to the new folder
        for doc_file in word_docs:
            old_path = os.path.join(folder_path, doc_file)
            new_path = os.path.join(word_folder_path, doc_file)
            shutil.move(old_path, new_path)

        print("Word documents have been organized in the folder:", word_folder_path)
    else:
        print("No Word documents found in the folder.")


# function is used to move all the word docs already present in the input folder in new folder AllDocuments
def move_word_documents(input_folder, output_folder):
    # Check if the "AllDocuments" folder exists, create it if not
    all_documents_folder = os.path.join(input_folder, "AllDocuments")
    if not os.path.exists(all_documents_folder):
        os.makedirs(all_documents_folder)

    # Get a list of all files in the input folder
    files = [file for file in os.listdir(input_folder) if file.endswith('.docx')]

    # Move each Word document to the "AllDocuments" folder
    for file in files:
        file_path = os.path.join(input_folder, file)
        destination_path = os.path.join(all_documents_folder, file)
        shutil.move(file_path, destination_path)
        print(f"Moved: {file} to {all_documents_folder}")


def excel_popup(windowName):
    while(True):
        #time.sleep(5)
        window_title = pgw.getActiveWindowTitle()
        global stop_threads
        logging.info('window_title ',window_title)

        logging.info(windowName)
        excel_windows = pgw.getWindowsWithTitle("Excel")
        try:
            for each_excel_window in excel_windows:
                logging.info("Title ", each_excel_window.title)
                # time.sleep(4)
                if (windowName.split('.')[0] in each_excel_window.title):

                    each_excel_window.minimize()
                    each_excel_window.maximize()
                    if (each_excel_window.isActive == False):
                        each_excel_window.activate()
                        break
                else:
                    each_excel_window.minimize()

        except Exception as e:
            # print("Exceptiion in Popup")
            logging.info("Exception in excel popup",e)
            pass
        if (pgw.getActiveWindowTitle() == "Microsoft Excel"):
            time.sleep(1)
            KMS.rightArrow()
            time.sleep(1)
            KMS.pressEnter()
        active_window = pgw.getActiveWindow()
        if active_window is not None:
            active_window.minimize()
        if stop_threads:
            break


def getPrevDocs(tpBook, vPT):
    combined_list = []
    logging.info("In getPrevDocs function = ",  vPT)
    rowx = 6
    ipDocList = []
    refList = []
    verList = []
    tpBook.activate()
    sheet = tpBook.sheets['Sommaire']
    tpBook.sheets['Sommaire'].activate()
    maxrow = tpBook.sheets['Sommaire'].range('A' + str(tpBook.sheets['Sommaire'].cells.last_cell.row)).end('up').row
    logging.info("maxrow-------------->", maxrow)
    try:
        logging.info("rowx < maxrow-------------->", rowx, maxrow)
        while rowx < maxrow + 1:
            if sheet.range(rowx, 1).value is not None:
                logging.info(sheet.range(rowx, 1).value, str(int(float(sheet.range(rowx, 1).value))), vPT)
                # if str(vPT) in str(sheet.range(rowx, 1).value):
                if str(vPT) == str(sheet.range(rowx, 1).value):
                    logging.info("HHere")
                    rlo = sheet.range(rowx, 1).merge_area.row
                    rhi = sheet.range(rowx, 1).merge_area.last_cell.row
                    logging.info("I---->", rlo, rhi)

                    greppedTypeIndex = []
                    greppedVersionIndex = []

                    # for i in range(rlo, rhi):
                    for i in range(rlo, rhi + 1):
                        logging.info(i)

                        if i not in greppedTypeIndex:
                            if sheet.range(i, 5).merge_cells is False:
                                if sheet.range(i, 5).value is not None:
                                    ipDocList.append(sheet.range(i, 5).value)

                            elif sheet.range(i, 5).merge_cells is True:
                                rlo = sheet.range(i, 5).merge_area.row
                                rhi = sheet.range(i, 5).merge_area.last_cell.row
                                mergedCount = rhi - rlo
                                logging.info("Values of rlo,rhi,mergedcount1--->", rlo, rhi, mergedCount)
                                logging.info("sheet.range(i, 5).value----------->", sheet.range(i, 5).value)
                                if sheet.range(i, 5).value is None:
                                    sheet.range(i, 5).unmerge()
                                    pass
                                if sheet.range(i, 5).value is not None:
                                    ipDocList.append(sheet.range(i, 5).value)

                                    for h in range(1, mergedCount + 1):
                                        logging.info("value of k--->", h)
                                        if sheet.range(i + h, 5).value is None:
                                            logging.info("value of i+k--->", sheet.range(i + h, 5).value)
                                            sheet.range(i + h, 5).value = sheet.range(i, 5).value
                                            ipDocList.append(sheet.range(i + h, 5).value)
                                            greppedTypeIndex.append(i + h)

                        if sheet.range(i, 6).merge_cells is False:
                            if sheet.range(i, 6).value is not None:
                                refList.append(sheet.range(i, 6).value)

                        elif sheet.range(i, 6).merge_cells is True:
                            rlo = sheet.range(i, 6).merge_area.row
                            rhi = sheet.range(i, 6).merge_area.last_cell.row
                            mergedCount = rhi - rlo
                            logging.info("Values of rlo,rhi,mergedcount2--->", rlo, rhi, mergedCount)
                            logging.info("sheet.range(i, 66).value----------->", sheet.range(i, 6).value)

                            if sheet.range(i, 6).value is None:
                                pass
                                # Additional actions if necessary

                            # sheet.range(i, 6).unmerge()
                            # tpBook.sheets['Sommaire'].range(i, 6).unmerge()
                            if sheet.range(i, 6).value is not None:
                                refList.append(sheet.range(i, 6).value)

                                for k in range(1, mergedCount + 1):
                                    logging.info("value of k--->", k)
                                    if sheet.range(i + k, 6).value is None:
                                        logging.info("value of i+k--->", sheet.range(i + k, 6).value)
                                        sheet.range(i + k, 6).value = sheet.range(i, 6).value
                                        refList.append(sheet.range(i + k, 6).value)

                        if i not in greppedVersionIndex:
                            if sheet.range(i, 7).merge_cells is False:
                                if sheet.range(i, 7).value is not None:
                                    version = re.search('[0-9]{1,2}', str(sheet.range(i, 7).value))
                                    if version is not None:
                                        # print("adding Version to version list1 =" + str(version) + "i = " + str(i))
                                        logging.info("adding Version to version list1  and value of i", version, i)
                                        verList.append(version.group())
                                    else:
                                        verList.append('')
                            elif sheet.range(i, 7).merge_cells is True:
                                rlo = sheet.range(i, 7).merge_area.row
                                rhi = sheet.range(i, 7).merge_area.last_cell.row
                                mergedCount = rhi - rlo
                                logging.info("Values of rlo,rhi,mergedcount3--->", rlo, rhi, mergedCount)
                                logging.info("sheet.range(i, 77).value----------->", sheet.range(i, 7).value)
                                # sheet.range(i, 7).unmerge()
                                if sheet.range(i, 7).value is None:
                                    pass

                                # tpBook.sheets['Sommaire'].range(i, 7).unmerge()
                                if sheet.range(i, 7).value is not None:
                                    version = re.search('[0-9]{1,2}', str(sheet.range(i, 7).value))
                                    if version is not None:
                                        # print("adding Version to version list2 =", + str(version) + "i = " + str(i) )
                                        logging.info("adding Version to version list2  and value of i", version, i)
                                        verList.append(version.group())
                                    else:
                                        verList.append('')

                                    for k in range(1, mergedCount + 1):
                                        logging.info("value of k--->", k)
                                        if sheet.range(i + k, 7).value is None:
                                            logging.info("value of i+k--->", sheet.range(i + k, 7).value)
                                            sheet.range(i + k, 7).value = sheet.range(i, 7).value
                                            logging.info("")
                                            version = re.search('[0-9]{1,2}', str(sheet.range(i + k, 7).value))
                                            if version is not None:
                                                # print("adding Version to version list3 =" + str(version) + "i = " + str(i))
                                                logging.info("adding Version to version list3  and value of i", version, i)
                                                verList.append(version.group())
                                                logging.info("adding i+k to greppedVersionIndex", i + k)
                                                greppedVersionIndex.append(i + k)
                                            else:
                                                verList.append('')
                        else:
                            logging.info("Already added the version to list", i)

            rowx = rowx + 1
        logging.info("ipDoc", ipDocList)
        logging.info("RefList", refList)
        logging.info("VerList", verList)
        logging.info("length------>", len(refList), len(verList))
        combined_list = list(zip(ipDocList, refList, verList))
        logging.info("combined_list------------>", combined_list)

        # Define the document types you want to check for
        document_types = ["ssd", "eead", "nt", "tfd"]

        # Initialize lists to store filtered data
        filtered_ipDoc = []
        filtered_RefList = []
        filtered_VerList = []

        # Iterate through each item in ipDoc and check if it contains the specified prefixes
        for i, doc in enumerate(ipDocList):
            for doc_type in document_types:
                if doc is not None and re.match(doc_type, doc.lower()):
                # if re.search(doc_type, doc.lower()):
                    # If the document type matches, add corresponding items to the filtered lists
                    filtered_ipDoc.append(doc)
                    filtered_RefList.append(refList[i])
                    filtered_VerList.append(verList[i])
                    break  # Break the inner loop as we found a match

        # Print the filtered lists (you can perform further operations as needed)
        logging.info("Filtered ipDoc:", filtered_ipDoc)
        logging.info("Filtered RefList:", filtered_RefList)
        logging.info("Filtered VerList:", filtered_VerList)
        combined_list = list(zip(filtered_RefList, filtered_VerList))
        logging.info("combined_list------------>", combined_list)
        logging.info("length of combined_list------------>", len(combined_list))

    except Exception as ex:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        exp_fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
        print(f"\nSomething went wrong {ex} line no. {exc_tb.tb_lineno} file name: {exp_fname}")
    return combined_list


def gettable_from_ssd_docs(table_lists, Doc_lists, req_not_present_in_download_docs_unique, req_present_in_download_docs):
    req_not_present_in_download_docs_unique = [item for item in req_not_present_in_download_docs_unique if item]
    for filename in os.listdir(ICF.getSsdFolder()):
        if filename.endswith(('.doc', '.docx')):
            # Process the file (replace this with your logic)
            logging.info(f"Processing getSsdFolder::: {filename}")
            if req_not_present_in_download_docs_unique[:]:
                logging.info("req_not_present_in_download_docs_unique------------>", req_not_present_in_download_docs_unique[:])
                try:
                    WDI.save_as_docx(ICF.getSsdFolder() + "\\" + filename)
                except:
                    pass
                Doc = os.path.join(ICF.getSsdFolder(), filename)
                logging.info("Keyword_file_path----->", Doc)
                for rv in req_not_present_in_download_docs_unique[:]:  # Iterate over a copy of the list
                    ReqName, ReqVer = getReqVer(rv)
                    print("reqName, reqVer---->", ReqName, ReqVer)
                    try:
                        RqTable, Content, Doc = WDI.getReqContent(Doc, ReqName, ReqVer)
                        if Content == -1:
                            print("Requirement not found in the file----------------->", ReqName, ReqVer, Content)
                        table_lists.append(RqTable)
                        Doc_lists.append(Doc)
                        if Content != None and Content != -1 and Content != '':
                            req_present_in_download_docs.append(rv)
                            req_not_present_in_download_docs_unique.remove(rv)  # Remove the element from the list
                    except Exception as ex:
                        exc_type, exc_obj, exc_tb = sys.exc_info()
                        exp_fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
                        print(
                            f"\nSomething went wrong getting table {ex} line no. {exc_tb.tb_lineno} file name: {exp_fname}")

            logging.info("req_present_in_download_docs------------>", req_present_in_download_docs)
            logging.info("rv_in_download_docs------------>", req_not_present_in_download_docs_unique)
    return req_not_present_in_download_docs_unique[:]


def main(sheets=None, reqs=None):
    print("entered supporting doc download")
    tpBook = EI.openTestPlan()
    all_refList_verList_tuples = []

    table_lists = []
    Doc_lists = []
    all_filtered_req_list = []
    req_present_in_download_docs = []
    try:
        if sheets:
            for sheet in tpBook.sheets:
                all_sheets = sheets.split(',')
                if sheet.name in all_sheets:
                    sheet.activate()
                    getReqList = sheet.range('C4').value.split("|")
                    logging.info("before req list = ", getReqList)
                    # Use list comprehension to remove empty elements
                    output_list = [item for item in getReqList if item]
                    # Print the filtered list
                    logging.info("after req list = ", output_list)
                    # after req list =  ['GEN-VHL-DC-BV-1210(2)', 'REQ-0229835(F)', 'REQ-0229816(C)', 'REQ-0533801(B)', 'GEN-VHL-DCINT-PUCSM.0049(7)', 'GEN-VHL-DCINT-PUCSM.0633(3)', 'GEN-VHL-DCINT-PUCSM.0504(6)', 'GEN-VHL-DCINT-PUCSM.0260(7)',
                    # 'GEN-VHL-DCINT-PUCSM.0094(15)', 'EV-VHL-NTINT.PUCSM.0110(0)', 'REQ-0679822(A)', 'GEN-VHL-DCINT-CVMM.0147(8)']
                    filtered_req_list = [req for req in output_list if req.startswith('REQ')]
                    logging.info("filtered_req_list------>", filtered_req_list)
                    all_filtered_req_list.extend(filtered_req_list)
                # filtered_req_list = ['REQ-0229835(F)', 'REQ-0229816(C)', 'REQ-0533801(B)', 'REQ-0679822(A)']
        elif reqs:
            all_filtered_req_list=reqs.split(',')
        all_filtered_req_list=[*set(all_filtered_req_list)]
        all_filtered_req_list_copy = all_filtered_req_list.copy()
        foundReq = []
        logging.info("all_filtered_req_list--------------->", all_filtered_req_list)
        logging.info("length of all_filtered_req_list--------------->", len(all_filtered_req_list))
        all_docs = getdocs(all_filtered_req_list, foundReq)
        notFoundReq = [*set(set(all_filtered_req_list)-set(foundReq))]

        if all_docs != -1:
            for filename in os.listdir(ICF.getInputFolder()):
                if filename.endswith(('.doc', '.docx')):
                    # Process the file (replace this with your logic)
                    logging.info(f"Processing getInputFolder::: {filename}")
                    try:
                        WDI.save_as_docx(ICF.getInputFolder() + "\\" + filename)
                    except:
                        pass
                    Doc = os.path.join(ICF.getInputFolder(), filename)
                    logging.info("Keyword_file_path----->", Doc)
                    for rv in all_filtered_req_list[:]:  # Iterate over a copy of the list
                        ReqName, ReqVer = getReqVer(rv)
                        logging.info("reqName, reqVer---->", ReqName, ReqVer)
                        try:
                            RqTable, Content, Doc = WDI.getReqContent(Doc, ReqName, ReqVer)
                            if Content == -1:
                                logging.info("Content----------------->", ReqName, ReqVer, Content)
                            table_lists.append(RqTable)
                            Doc_lists.append(Doc)
                            if Content != None and Content != -1 and Content != '':
                                req_present_in_download_docs.append(rv)
                                all_filtered_req_list.remove(rv)  # Remove the element from the list
                        except Exception as ex:
                            exc_type, exc_obj, exc_tb = sys.exc_info()
                            exp_fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
                            print(
                                f"\nSomething went wrong getting table {ex} line no. {exc_tb.tb_lineno} file name: {exp_fname}")

                    logging.info("req_present_in_download_docs------------>", req_present_in_download_docs)
                    logging.info("rv_in_download_docs------------>", all_filtered_req_list)

        req_present_in_download_docs_unique = [*set(req_present_in_download_docs)]
        req_not_present_in_download_docs_unique = [*set(all_filtered_req_list)]
        vPT = []

        for req in req_not_present_in_download_docs_unique:
                for sheet in tpBook.sheets:

                    if sheets:
                        all_sheets = sheets.split(',')
                        if sheet.name in all_sheets:
                            flag = True
                        else:
                            flag = False
                    else:
                        flag = True
                    if flag and ('0000' not in sheet.name) and ('VSM' in sheet.name) and sheet.visible:
                        sheet.activate()
                        getReqList = sheet.range('C4').value.split("|")
                        logging.info("before req list = ", getReqList)
                        # Use list comprehension to remove empty elements
                        output_list = [item for item in getReqList if item]
                        if req in output_list:
                            result = EI.searchDataInCol(sheet,3,'Nature des modifications')
                            value = EI.getDataFromCell(sheet,(result['cellPositions'][0][0]+1,result['cellPositions'][0][1]-1))
                            count = 2
                            while(value is not None):
                                vPT.append(value)
                                value = EI.getDataFromCell(sheet, (
                                result['cellPositions'][0][0] + count, result['cellPositions'][0][1] - 1))
                                count = count+1

        print(f"{len([*set(foundReq)])} Requirements found in analyse de entrant---->{[*set(foundReq)]}")
        print(f"{len(notFoundReq)} Requirements not found in analyse de entrant---->{notFoundReq}")
        print(
            f"{len(req_present_in_download_docs_unique)} Requirements found in analyse de entrant---->{req_present_in_download_docs_unique}")
        print(
            f"{len(req_not_present_in_download_docs_unique)} Requirements not found in analyse de entrant---->{req_not_present_in_download_docs_unique}")

        vPT = [*set(vPT)]
        print("vPT----->", vPT)
        ssd_folder = ICF.getSsdFolder()
        if not os.path.exists(ssd_folder) or not os.listdir(ssd_folder):
            # Either the directory doesn't exist or it's empty
            # Handle the condition here
            logging.info("Directory is either not present or empty.")
            TestPlan = EI.findInputFiles()[1]
            logging.info("PT---->", TestPlan)
            tpBook = EI.openExcel(ICF.getInputFolder() + "\\" + TestPlan)
            combined_list = []
            for VPT in vPT:
                result = getPrevDocs(tpBook, VPT)  # Call the getPrevDocs function
                combined_list.extend(result)
            logging.info(combined_list)
            combined_list = list(set(combined_list))
            print("duplictes remove from combined_list------------->", combined_list)

            filtered_list = [pair for pair in combined_list if pair != (None, '')]
            print("filtered_list--------->", filtered_list)

            for referenceNum, VersionNum in filtered_list:
                SR.startDocumentDownload([[referenceNum, VersionNum]], True, True, False)
        req_not_present_in_download_docs_unique = gettable_from_ssd_docs(table_lists, Doc_lists, req_not_present_in_download_docs_unique, req_present_in_download_docs)

        logging.info("before_table_lists--------->", table_lists)
        logging.info("before_Doc_lists--------->", Doc_lists)

        result_tuples = list(zip(table_lists, Doc_lists))
        logging.info("before_result_tuples---->", result_tuples)
        table_lists = [(a, b) for (a, b) in result_tuples if a != '' and a != -1]
        logging.info("before_result_tuples---->", table_lists)

        if req_not_present_in_download_docs_unique:
            print(
                f"{len(req_not_present_in_download_docs_unique)} Requirements not found in analyse de entrant and SSD documents---->{', '.join(req_not_present_in_download_docs_unique)}. Go for the Global Search.")
        else:
            print("All reqs are present in the output document.")

        for data in table_lists:
            addDataInDocument(data, "T")
        print("Processing complete.")
        tpBook.close()
    except Exception as ex:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        exp_fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
        print(f"\nSomething went wrong {ex} line no. {exc_tb.tb_lineno} file name: {exp_fname}")

# VSM20_N1_20_77_0001,VSM20_N1_20_77_0018E,VSM20_GC_20_77_0038A,VSM20_GC_20_77_0024B,VSM20_N1_20_77_0025A
# VSM20_N1_20_47_0011,VSM20_N1_20_47_0001


if __name__ == '__main__':
    start = time.time()
    print("Tool start time---------->", start)
    ICF.loadConfig()
    choice = input("Enter 1 for sheets or 2 for requirements : ")
    if choice == '1':
        sheets = input("Enter the sheets names by giving , : ")
        main(sheets)
    elif choice == '2':
        reqs = input("Enter the requirement names by giving , : ")
        main(None, reqs)
    end1 = time.time()
    print("\nexecution time " + str(end1 - start))