import os
import re
import sys
import time
import docx
import shutil
import logging
import datetime
import xlwings as xw
from docx import Document
from selenium import webdriver
from docx.enum.text import WD_BREAK
import InputConfigParser as ICF
from selenium.common import exceptions
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as Exp_Con

date_time = datetime.datetime.now()


def openExcel(book: any) -> any:
    return xw.Book(book)


def getDataFromCell(sheet, colRow) -> any:
    return sheet.range(colRow).value


def open_TestPlan_SearchLogic() -> any:
    PT = ''
    if not os.path.exists(ICF.getInputFolder()):
        os.makedirs(ICF.getInputFolder())
    arr = os.listdir(ICF.getInputFolder())
    for i in arr:
        if i.find('Tests')!=-1 and i.find('fiches')==-1 and i.find('~$')==-1:
            PT = i
    if len(PT)!=0:
        testPlan = openExcel(ICF.getInputFolder() + "\\" + PT)
        return testPlan
    else:
        print("Testplan is not present in the input folder")
        exit(404)


def oldFormatContent(doc, Keyword) -> dict:
    searchResult = {}
    for table in doc.tables:
        num_rows = len(table.rows)
        num_columns = len(table.columns)
        if num_columns==3 and num_rows >= 6:
            req_id_cell = table.rows[5].cells[0]
            req_id = req_id_cell.text.strip()
            logging.info('Keyword---->', Keyword, "req_id------>", req_id)
            if Keyword in req_id:
                logging.info('keyword in content')
                searchResult.update({"reqId": table.rows[0].cells[0].text.strip()})
                searchResult.update({"table": table})
                for row_index, row in enumerate(table.rows):
                    try:
                        if "content of the requirement" in row.cells[0].text.lower().strip():
                            next_row = table.rows[row_index + 1]
                            content_cell = next_row.cells[0]
                            clearStrikethrough(content_cell)
                            searchResult.update({"content": str(content_cell.text)})
                        if "Effectivity" in row.cells[0].text.strip():
                            next_row = table.rows[row_index + 1]
                            effectivity_cell = next_row.cells[0]
                            clearStrikethrough(effectivity_cell)
                            searchResult.update({"effectivity": str(effectivity_cell.encode('utf-8').strip())})
                        elif "LCDV" in row.cells[0].text.strip():
                            next_row = table.rows[row_index + 1]
                            lcdv_cell = next_row.cells[0]
                            clearStrikethrough(lcdv_cell)
                            searchResult.update({"LCDV": str(lcdv_cell.text.encode('utf-8').strip())})
                        elif "diversity" in row.cells[0].text.lower().strip():
                            next_row = table.rows[row_index + 1]
                            print('next_row.cells.text[0]----->', next_row.cells[0].text)
                            diversity_cell = next_row.cells[0]
                            clearStrikethrough(diversity_cell)
                            searchResult.update({"diversity": str(diversity_cell.text.encode('utf-8').strip())})
                        elif "target configuration" in row.cells[0].text.lower().strip():
                            next_row = table.rows[row_index + 1]
                            target_cell = next_row.cells[1]
                            clearStrikethrough(target_cell)
                            searchResult.update({"target": str(target_cell.text.encode('utf-8').strip())})
                    except Exception as exp:
                        exc_type, exc_obj, exc_tb = sys.exc_info()
                        exp_fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
                        print(
                            f"\nProblem in fetching the content from old format document {exp} line no. {exc_tb.tb_lineno} file name: {exp_fname}********************")
    return searchResult


def newFormatContent(doc, Keyword) -> dict:
    searchResult = {}
    for table in doc.tables:
        num_rows = len(table.rows)
        num_columns = len(table.columns)
        if num_columns==2 and num_rows >= 5:
            req_id_cell = table.rows[1].cells[1]
            req_id = req_id_cell.text.strip()
            if Keyword in req_id:
                searchResult.update({"table": table})
                try:
                    content_cell = table.rows[1].cells[1]
                    clearStrikethrough(content_cell)
                    searchResult.update({"reqId": table.rows[1].cells[0].text.strip(), "content": content_cell.text})
                except Exception as e:
                    print("Exception in getOldContents = ", e)
                    searchResult.update(
                        {"reqId": table.rows[1].cells[0].text.strip(), "content": table.rows[1].cells[1].text})
                for row in table.rows:
                    try:
                        if "Effectivity" in row.cells[0].text.strip():
                            searchResult.update(
                                {"effectivity": str(row.cells[1].text.encode('utf-8').strip())})
                        elif "LCDV" in row.cells[0].text.strip():
                            searchResult.update(
                                {"LCDV": str(row.cells[1].text.encode('utf-8').strip())})
                        elif "diversity" in row.cells[0].text.lower().strip():
                            searchResult.update(
                                {"diversity": str(row.cells[1].text.encode('utf-8').strip())})
                        elif "target configuration" in row.cells[0].text.lower().strip():
                            searchResult.update(
                                {"target": str(row.cells[1].text.encode('utf-8').strip())})
                    except Exception as exp:
                        exc_type, exc_obj, exc_tb = sys.exc_info()
                        exp_fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
                        print(
                            f"\nProblem in fetching the content from new format document {exp} line no. {exc_tb.tb_lineno} file name: {exp_fname}********************")
    return searchResult


def find_requirement_content(file_path, Keyword) -> dict:
    print(f"finding Signal content {Keyword}")
    searchResult = {}
    doc = docx.Document(file_path)
    num_tables = len(doc.tables)
    logging.info("Total Number of Tables = ", num_tables)
    req_res = newFormatContent(doc, Keyword)
    logging.info(f"\n>>>>req_res new: {req_res}")
    if req_res:
        searchResult = req_res
    else:
        logging.info("searching data in old format...")
        req_res = oldFormatContent(doc, Keyword)
        logging.info(f"\n>>>>req_res old: {req_res}")
        if req_res:
            searchResult = req_res
    return searchResult


def clearStrikethrough(cell) -> None:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            # Check whether the run text is strikethrough
            if run.font.strike:
                # Remove the strikethrough text
                run.clear()


def docInfo(refnum, supporting_req: bool, Global_search: bool, reqVer="") -> None:
    username, password = ICF.getCredentials()
    driver.get(
        'https://' + username + ':' + password + '@docinfogroupe.psa-peugeot-citroen.com/ead/accueil_init.action')
    wait = WebDriverWait(driver, 10)
    wait.until(lambda driver: driver.execute_script('return document.readyState')=='complete')
    expand_menu = driver.find_element(By.ID, "ext-gen24")
    time.sleep(1)
    expand_menu.click()
    time.sleep(1)
    ref_search = driver.find_element(By.ID, "txtRef")
    ref_search.click()
    time.sleep(1)
    ref_search.send_keys(refnum)
    time.sleep(1)
    ref_search.send_keys(Keys.RETURN)
    wait.until(lambda driver: driver.execute_script('return document.readyState')=='complete')
    content = driver.find_elements(By.CLASS_NAME, "x-tab-strip-text")
    for i in content:
        if i.text=="Versions":
            logging.info(i.text)
            i.click()
    time.sleep(5)
    try:
        if reqVer!="":
            doc_download(refnum, reqVer, supporting_req, Global_search)
        else:
            lst_ver = []
            version_element = driver.find_elements(By.XPATH, (
                "//div[@class='x-grid3-cell-inner x-grid3-col-docVersion']"))  # version table
            for x in version_element:
                actualVer = (float(x.get_attribute('innerHTML')))
                lst_ver.append(actualVer)
            logging.info(lst_ver)
            ver_num = []
            for num in lst_ver:
                val_ver = str(num).split('.')
                if int(val_ver[1])==0:
                    ver_num.append(float(val_ver[0]))
            doc_download(refnum, max(ver_num), supporting_req, Global_search)
            logging.info(max(ver_num))
    except Exception as ex:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        print(f"Exception in Downloading the Documents.{ex}{exc_tb.tb_lineno}")


def doc_download(refnum, reqVer, supporting_req, Global_search) -> None:
    table_body = driver.find_elements(By.CLASS_NAME, "x-grid3-scroller")  # slide bar
    wait = WebDriverWait(driver, 15)
    logging.info("table_body------>", table_body)
    for tbody in table_body:
        if not tbody.find_elements(By.CLASS_NAME,
                                   "x-grid3-col x-grid3-cell x-grid3-td-docViewIcon x-grid3-cell-first "):
            tables = tbody.find_elements(By.TAG_NAME, "table")
            for each_table in tables:
                try:
                    t_body = each_table.find_element(By.TAG_NAME, "tbody")
                    each_row = t_body.find_element(By.TAG_NAME, "tr")
                    all_data = each_row.find_elements(By.TAG_NAME, "td")
                    lst_elements = []
                    txt_elements = []
                    for each_data in all_data:
                        # time.sleep(2)
                        ver_txt = each_data
                        lst_elements.append(ver_txt)
                        txt_elements.append(ver_txt.text)
                    logging.info(txt_elements)
                    if len(lst_elements) >= 3:
                        if str(float(reqVer))==lst_elements[0].text:
                            logging.info("Version matched")
                            url = lst_elements[2].find_element(By.TAG_NAME, "a")
                            a = Exp_Con.element_to_be_clickable(url)
                            wait.until(a).click()
                            count = 0
                            attachment = wait.until(Exp_Con.element_to_be_clickable((By.ID, "attachementPanel")))
                            attached_elements = attachment.find_elements(By.TAG_NAME, ('a'))
                            toggle = attachment.find_elements(By.XPATH, "//img[@title='Click to collapse.']")
                            words = ["doc"]
                            for attached in attached_elements[0:]:
                                if attached not in toggle:
                                    logging.info('file_name', attached.text)
                                    for w in words:
                                        if w in attached.text:
                                            count = count + 1
                                            attached.click()
                            sortFiles(count, reqVer, refnum, supporting_req, Global_search)
                            driver.close()
                            break
                        logging.info("Version not matched")
                except exceptions.StaleElementReferenceException as e:
                    print(e)


def moveFile(x: any, supporting_req: bool, Global_search: bool) -> None:
    logging.info('supporting_req---->', supporting_req)
    logging.info('Global_search---->', Global_search)
    guest = ICF.getSsdFolder()
    if supporting_req is True and Global_search is False:
        dst = guest
        if not os.path.exists(dst):
            os.makedirs(dst)
        logging.info('dst---0->', dst, "supporting_req--->", supporting_req)
    elif Global_search is True and supporting_req is False:
        dst = guest + "\\" + "Buffer"
        if not os.path.exists(dst):
            os.makedirs(dst)
        logging.info('dst---1->', dst, "Global_search--->", Global_search)
    else:
        dst = ICF.getSsdFolder()
        logging.info('dst---2->', dst)
    src = ICF.getDownloadFolder()
    sorc = (os.path.join(src, x))
    dset = (os.path.join(dst, x))
    logging.info('src====>', sorc)
    logging.info('dst====>', dset)
    logging.info("Before moving file:")
    logging.info(os.listdir(dst))
    if os.path.isfile(dset):
        os.remove(dset)
        logging.info(x, 'deleted in', dst)
    dest = shutil.move(sorc, dset)
    logging.info('moved file:', dest)
    logging.info('file Moved successfully')


def sortFiles(count, reqVer, refnum, supporting_req: bool, Global_search: bool) -> int:
    while any(filename.endswith('.crdownload') or filename.endswith('.tmp') for filename in
              os.listdir(ICF.getDownloadFolder())):
        time.sleep(1)
    a = -1
    path = ICF.getDownloadFolder()
    dst = ICF.getSsdFolder()
    name_list = os.listdir(path)
    full_list = [os.path.join(path, i) for i in name_list]
    time_sorted_list = sorted(full_list, key=os.path.getmtime)
    sorted_filename_list = [os.path.basename(i) for i in time_sorted_list]
    sorted_filename_list.reverse()
    for x in (sorted_filename_list[0:count]):
        if (os.path.splitext(x)[1]==".docx") or (os.path.splitext(x)[1]==".doc") or (
                os.path.splitext(x)[1]==".docm") or (os.path.splitext(x)[1]==".rtf"):
            if refnum not in x:
                logging.info(refnum)
                newFile = '[' + 'V' + str(float(reqVer)) + ']' + '[' + str(refnum) + ']' + x
            else:
                newFile = '[' + 'V' + str(float(reqVer)) + ']' + x
            logging.info("- ", newFile, path + '\\' + x, path + '\\' + newFile)
            try:
                os.rename(path + '\\' + x, path + '\\' + newFile)
            except:
                os.remove(path + '\\' + newFile)
                time.sleep(3)
                os.rename(path + '\\' + x, path + '\\' + newFile)
            logging.info("f------>", newFile)
            moveFile(newFile, supporting_req, Global_search)
            a = dst + '\\' + newFile
        else:
            logging.info("x----->", x)
            moveFile(x, supporting_req, Global_search)
    return a


def delete_files_with_extensions() -> None:
    downloads_path = os.path.join(os.path.expanduser("~"), "Downloads")
    extensions = [".crdownload", ".crdownloads", ".temp", ".tmp"]
    for root, _, files in os.walk(downloads_path):
        for file in files:
            _, file_extension = os.path.splitext(file)
            if file_extension.lower() in extensions:
                file_path = os.path.join(root, file)
                try:
                    os.remove(file_path)
                    print(f"Deleted: {file_path}")
                except Exception as e:
                    print(f"Error deleting {file_path}: {e}")


def startDocumentDownload(fepsRefVers, allowDownload: bool, supporting_req: bool, Global_search: bool) -> None:
    logging.info('ICF.getWebDriver()---->', os.path.exists(ICF.getWebDriver()))
    if os.path.exists(ICF.getWebDriver()):
        print("startDocumentDownload +", fepsRefVers)
        fepsRefVers = [item for item in fepsRefVers if item!=(None, '')]
        fepsRefVers = [[item[0].strip(), item[1]] for item in fepsRefVers]
        logging.info("after remove none elements--->", fepsRefVers)
        if allowDownload:
            logging.info("Allowed", allowDownload)
            for fepdDoc in fepsRefVers:
                logging.info("Start download = ", fepdDoc)
                ref, ver = fepdDoc
                if ref!="":
                    delete_files_with_extensions()
                    service = Service(executable_path=ICF.getWebDriver())
                    options = webdriver.ChromeOptions()
                    options.add_experimental_option('excludeSwitches', ['enable-logging'])
                    global driver
                    driver = webdriver.Chrome(service=service, options=options)
                    logging.info("--", ref, ver)
                    docInfo(ref, supporting_req, Global_search, ver.replace("V", ""))
                    time.sleep(3)
                else:
                    print("reference number is empty")
    else:
        print('Config Folder is not present')
        exit(404)


def getCellAbsVal(sheet, row, col) -> any:
    for i in range(row, 0, -1):
        cellVal = getDataFromCell(sheet, f"{col}{i}")
        if cellVal is not None:
            return cellVal
    return None


def downloadSSD(tpBook) -> None:
    summarySheet = any
    try:
        summarySheet = tpBook.sheets["Sommaire"]
    except Exception as e:
        print(f"TestPlan or Sommaire sheet not found! in Input Path.", e)
        exit(1)
    referenceList = []
    VersionList= []
    # Define the document types you want to check for
    document_types = ["ssd", "eead", "nt", "tfd"]
    # get all reference numbers for ssd files
    nrows = summarySheet.used_range.last_cell.row
    for i in range(6, nrows):
        typeVal = getDataFromCell(summarySheet, f"E{i}")
        for a in document_types:
            if typeVal is not None and re.match(a, typeVal.lower()):
                referenceNumber = getCellAbsVal(summarySheet, i, "F")
                versionNumber = getCellAbsVal(summarySheet, i, "G")
                referenceList.append(referenceNumber.strip())
                VersionList.append(versionNumber.strip())
    logging.info("referenceList --> ", referenceList, VersionList)
    print("referenceList, VerList--> ", len(referenceList), len(VersionList))
    combined_list = list(zip(referenceList, VersionList))
    print("combined_list------------>", combined_list)
    print("length of combined_list------------>", len(combined_list))
    unique_combined_list = list(set(combined_list))
    print("unique_combined_list------------>", unique_combined_list)
    print("length of unique_combined_list------------>", len(unique_combined_list))
    for referenceNum, VersionNum in unique_combined_list:
        startDocumentDownload([[referenceNum, VersionNum]], True, True, False)


def find_and_save_keyword(file_path, keyword, Keyword_file_folder) -> None:
    try:
        doc = Document(file_path)
        if any(keyword in paragraph.text for paragraph in doc.paragraphs):
            if not os.path.exists(Keyword_file_folder):
                os.makedirs(Keyword_file_folder)
            file_name = os.path.basename(file_path)
            output_path = os.path.join(Keyword_file_folder, file_name)
            doc.save(output_path)
            logging.info(f"Keyword found in '{file_name}'. Document saved to '{Keyword_file_folder}'.")
    except Exception as e:
        print(f"Error processing find_and_save_keyword func() -'{file_path}': {e}")


def process_folder(input_folder, keyword, Keyword_file_folder) -> None:
    files = [f for f in os.listdir(input_folder) if f.endswith(".docx")]
    for file in files:
        file_path = os.path.join(input_folder, file)
        find_and_save_keyword(file_path, keyword, Keyword_file_folder)


def getTables(fileName) -> any:
    wordDoc = docx.Document(fileName)
    return wordDoc.tables


def find_tables_by_keyword_and_header(table, keyword, header) -> list:
    matching_tables = []
    if table.cell(0, 0).text.strip().lower() == header.lower():
        for row in table.rows:
            for cell in row.cells:
                if keyword in cell.text:
                    matching_tables.append(table)
                    break
            if matching_tables:
                break
    return matching_tables


def search_Supporting_Req(input_folder, Keyword_file_folder, keyword):
    header = "N°Requirement"
    matching_tables = []
    flattened_list = ''
    logging.info('Keyword_file_folder---00->', Keyword_file_folder)
    if not os.path.exists(Keyword_file_folder):
        logging.info('Keyword_file_folder-11--->', Keyword_file_folder)
        os.makedirs(Keyword_file_folder)
    process_folder(input_folder, keyword, Keyword_file_folder)
    for filename in os.listdir(Keyword_file_folder):
        if filename.endswith(".docx"):
            Keyword_file_path = os.path.join(Keyword_file_folder, filename)
            if os.path.exists(Keyword_file_path):
                try:
                    logging.info('Keyword_file_path--->', Keyword_file_path)
                    requirement_content = find_requirement_content(Keyword_file_path, keyword)
                    logging.info('requirement_content in the doc--->', requirement_content)
                    if requirement_content:
                        logging.info('requirement_content["table"]---->', requirement_content["table"])
                        matching_tables.append((filename, requirement_content["table"]))
                    # flattened_list = [table for sublist in matching_tables for table in sublist]
                    logging.info("matching_tables---->", matching_tables)
                except Exception as ex:
                    exc_type, exc_obj, exc_tb = sys.exc_info()
                    print(f"Requirement is not in the correct format please Process Manual.{ex}{exc_tb.tb_lineno}")
    print('matching_tables---->', matching_tables)
    return matching_tables


def is_document_empty(doc):
    has_paragraphs = bool(doc.paragraphs)
    has_tables = bool(doc.tables)
    has_headings = any(p.style.name.startswith('Heading') for p in doc.paragraphs)
    return not (has_paragraphs or has_tables or has_headings)


def addDataInDocument(data, type) -> int:
    logging.info(f"\n-------> data <-------\n{data}")
    output_dir = os.path.abspath(r"..\Supporting_Requirement_Output")
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    if os.path.exists(r'..\Supporting_Requirement_Output\Supporting_Requirement.docx'):
        document = Document(r'..\Supporting_Requirement_Output\Supporting_Requirement.docx')
    else:
        document = Document()

    is_empty = is_document_empty(document)
    if not is_empty and type == 'H':
        run = document.add_paragraph().add_run()
        run.add_break(WD_BREAK.PAGE)
    if type == 'H':
        document.add_heading(f"--------------- {data[0]} -------------", level=1)
        document.add_heading(f"--------------- {data[1]} -------------", level=1)
    elif type == 'T':
        document.add_heading(f"--------------- {data[0]} -------------", level=1)
        document.add_paragraph('\n')
        destination_table = document.add_table(rows=len(data[1].rows), cols=len(data[1].columns))
        for row_idx, row in enumerate(data[1].rows):
            for col_idx, cell in enumerate(row.cells):
                destination_table.cell(row_idx, col_idx).text = cell.text
    else:
        document.add_paragraph(data[1])
    output_dir = os.path.abspath(r"..\Supporting_Requirement_Output")
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    savingPath = os.path.abspath(r'..\Supporting_Requirement_Output\Supporting_Requirement.docx')
    document.save(savingPath)
    return 1




if __name__ == "__main__":
    ICF.loadConfig()
    # main()
    tpBook = open_TestPlan_SearchLogic()
    end0 = time.time()
    # print("\n\n execution time for extracting data from Testplan " + str(end0 - start0))
    start1 = time.time()
    downloadSSD(tpBook)
    tpBook.close()
